"""
will_builder.py (final)
- _val_or_ph(): ดึงข้อมูลจริง / fallback placeholder สีแดง
- generate_clause3(): Claude gen + proofread
- build_will_docx(): ใช้ fullname_* / id_* จาก client_data
- build_will(): entry point
"""
import os, json, tempfile, zipfile, shutil, re
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT  = "Cordia New"
SZ    = 16
SZ_H  = 18
SZ_T  = 24
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x66, 0x66, 0x66)
RED   = RGBColor(0xAA, 0x00, 0x00)
FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")

_NA = ('ยังไม่มี', '-', '[ยังไม่ได้กำหนด]', 'ไม่ได้ระบุ', 'ยังไม่ได้กำหนด', '')


# ── helpers ────────────────────────────────────────────────────────────
def _clean(name: str) -> str:
    return re.sub(r'\s*\(.*?\)', '', name or '').strip()

def _ph(label: str, name: str) -> str:
    n = _clean(name)
    return f"[{label}]" if not n or n in _NA else f"[{label}{n}]"

def _val_or_ph(value: str, placeholder: str) -> str:
    """ถ้ามีข้อมูลจริง → คืนตรงๆ / ไม่มี → [placeholder] สีแดง"""
    v = (value or '').strip()
    return v if v and v not in _NA else f"[{placeholder}]"


# ── text rendering ─────────────────────────────────────────────────────
def _run(para, text, bold=False, size=None, color=None, underline=False):
    r = para.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size or SZ)
    r.bold = bold; r.font.color.rgb = color or BLACK
    if underline: r.underline = True
    return r

def _add_runs(para, text, size=None):
    for p in re.split(r'(\[.*?\])', text):
        if not p: continue
        if p.startswith('['):
            _run(para, p, bold=True, size=size, color=RED, underline=True)
        else:
            _run(para, p, size=size)

def _para(doc, text, size=None, bold=False, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt((size or SZ) * 1.5)
    p.alignment = align
    if bold or color:
        _run(p, text, bold=bold, size=size, color=color)
    else:
        _add_runs(p, text, size=size)

def _clause(doc, n, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(SZ_H * 1.5)
    _run(p, f"ข้อ {n}  {title}", bold=True, size=SZ_H)

def _subhead(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(SZ * 1.5)
    _run(p, text, bold=True, size=SZ)

def _spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(pts)
    p.paragraph_format.line_spacing = Pt(pts + 2)

def _sig_line(doc, role, name_ph):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(SZ * 1.5)
    _run(p, "ลงชื่อ  .......................................  ")
    _run(p, role)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(4)
    p2.paragraph_format.line_spacing = Pt(SZ * 1.5)
    _run(p2, "        (", size=SZ - 1, color=GRAY)
    _add_runs(p2, name_ph, size=SZ - 1)
    _run(p2, ")", size=SZ - 1, color=GRAY)


# ── Claude: generate + proofread ข้อ ๓ ────────────────────────────────
CLAUSE3_TEMPLATES = {
    "cash":       {"key": "เงินสดและบัญชีธนาคาร",
                   "text": "ยกเงินสด เงินในบัญชีธนาคารทุกบัญชี และเงินออมทุกประเภทที่มีในชื่อข้าพเจ้า ณ วันที่ข้าพเจ้าถึงแก่ความตาย ให้แก่ {recipient} {relationship}"},
    "investment": {"key": "หุ้นและกองทุนรวม",
                   "text": "ยกหุ้นสามัญในตลาดหลักทรัพย์แห่งประเทศไทย กองทุนรวม และกองทุนสำรองเลี้ยงชีพทั้งหมดที่มีในชื่อข้าพเจ้า ให้แก่ {recipient} โดยให้โอนเข้าชื่อบุตรแต่ละคนตามกฎหมาย"},
    "property":   {"key": "อสังหาริมทรัพย์",
                   "text": "ยกบ้านพร้อมที่ดิน โฉนดเลขที่ {deed_no} ตั้งอยู่ที่ {address} ให้แก่ {recipient} โดยให้โอนเข้าชื่อบุตรแต่ละคนตามกฎหมาย"},
    "crypto":     {"key": "สินทรัพย์ดิจิทัล",
                   "text": "ยกสินทรัพย์ดิจิทัลทั้งหมดในกระเป๋าอิเล็กทรอนิกส์ที่มีในชื่อข้าพเจ้า ให้แก่ {recipient} โดยข้อมูลการเข้าถึงได้จัดเก็บไว้ที่ {storage} ผู้จัดการมรดกสามารถเปิดได้เมื่อข้าพเจ้าถึงแก่ความตายเท่านั้น"},
    "business":   {"key": "กิจการและหุ้นส่วน",
                   "text": "ยกหุ้นส่วนและสิทธิ์ทั้งหมดในกิจการ {business_name} ให้แก่ {recipient}"},
    "valuables":  {"key": "ทองคำและทรัพย์สินมีค่า",
                   "text": "ยกทองคำ เครื่องประดับ และทรัพย์สินมีค่าทั้งหมด ให้แก่ {recipient} {relationship}"},
    "other":      {"key": "ทรัพย์สินอื่นๆ",
                   "text": "ทรัพย์สินอื่นๆ ที่มิได้ระบุไว้ข้างต้น ให้ตกเป็นของ {recipient} {relationship}"},
}

def generate_clause3(client_data: dict) -> list:
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    fullname_spouse   = client_data.get("fullname_spouse", "")
    fullname_children = client_data.get("fullname_children", "")

    prompt = f"""คุณคือผู้ช่วยร่างพินัยกรรมของคุณพยัต นักวางแผนกฎหมายและการเงินในประเทศไทย

ข้อมูลลูกค้า:
{json.dumps(client_data, ensure_ascii=False, indent=2)}

Templates ที่มีให้เลือก:
{json.dumps(CLAUSE3_TEMPLATES, ensure_ascii=False, indent=2)}

งาน: เขียน sub-clause ข้อ ๓ การแบ่งทรัพย์สิน

กฎ:
1. เลือก template เฉพาะทรัพย์สินที่ลูกค้ามีจริงเท่านั้น (ดูจาก assets_* และ แผนแบ่งทรัพย์)
2. template "other" ต้องมีเสมอ วางไว้ท้ายสุด
3. เรียงลำดับ ๓.๑ ๓.๒ ๓.๓ ... ตามลำดับ
4. ระบุผู้รับมรดกจาก field แผนแบ่งทรัพย์ของลูกค้า
5. ถ้ามีชื่อ-นามสกุลจริงของคู่สมรส ให้ใช้ชื่อจริง: "{fullname_spouse}"
6. ถ้ามีชื่อ-นามสกุลจริงของลูก ให้ใช้ชื่อจริง: "{fullname_children}" (แยกด้วย ,)
7. ถ้าไม่มีชื่อจริง ให้ใช้ placeholder เช่น [ชื่อ-นามสกุลคุณจูน] หรือ [ชื่อ-นามสกุลน้องฟ้า]
8. ถ้าไม่รู้โฉนด/ที่อยู่ที่ดิน ให้ใส่ [เลขโฉนดที่ดิน] และ [ที่อยู่บ้านและที่ดิน]
9. ถ้าไม่รู้ที่เก็บ crypto ให้ใส่ [ที่เก็บข้อมูลการเข้าถึง]
10. ใช้ภาษาไทยล้วน ไม่มีวงเล็บภาษาอังกฤษ
11. placeholder ใช้รูปแบบ [ข้อความอธิบาย] ห้ามใช้คำว่า "จริง" ใน placeholder
12. ใช้บริบทกฎหมายและสถาบันการเงินไทยเท่านั้น

output เป็น JSON เท่านั้น ไม่มีข้อความอื่น ไม่มี markdown:
[{{"num": "๓.๑", "title": "เงินสดและบัญชีธนาคาร", "text": "ยกเงินสด..."}}]"""

    r1 = client.messages.create(
        model="claude-sonnet-4-6", max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    text = r1.content[0].text.replace("```json","").replace("```","").strip()
    s = text.find("["); e = text.rfind("]") + 1
    result = json.loads(text[s:e])

    # ── proofread ──────────────────────────────────────────────────────
    proof_prompt = f"""ตรวจสอบและแก้ไข JSON นี้ให้ถูกต้อง

กฎ:
1. typo ภาษาไทย เช่น "ทรัพย์" ไม่ใช่ "ทรัพ์" "พร้อม" ไม่ใช่ "พร้วม"
2. ห้ามมีคำภาษาอังกฤษยกเว้น ชื่อเฉพาะที่จำเป็น
3. ภาษากฎหมายต้องถูกต้องและเป็นทางการ
4. placeholder ต้องอยู่ในรูป [ข้อความ] ครบถ้วน
5. ส่งกลับ JSON เหมือนเดิมทุกอย่าง แก้แค่จุดที่ผิด

{json.dumps(result, ensure_ascii=False)}"""

    r2 = client.messages.create(
        model="claude-sonnet-4-6", max_tokens=3000,
        messages=[{"role": "user", "content": proof_prompt}]
    )
    pt = r2.content[0].text.replace("```json","").replace("```","").strip()
    ps = pt.find("["); pe = pt.rfind("]") + 1
    if ps != -1 and pe > 0:
        result = json.loads(pt[ps:pe])
        print("✅ Proofread เสร็จ")
    else:
        print("⚠️ Proofread ไม่ได้ JSON ใช้ข้อมูลเดิม")

    return result


# ── build docx ─────────────────────────────────────────────────────────
def build_will_docx(client_data: dict, clause3_items: list) -> str:
    d = client_data
    n = d.get("ชื่อเล่น", d.get("nickname", "ผู้ทำพินัยกรรม"))

    # ── ดึงข้อมูลจริง + fallback ──────────────────────────────────────
    fs   = _val_or_ph(d.get("fullname_self", ""),   f"ชื่อ-นามสกุลคุณ{n}")
    ids  = _val_or_ph(d.get("id_self", ""),          f"เลขบัตรประชาชนคุณ{n}")
    addr = _val_or_ph(d.get("address_self", ""),     f"ที่อยู่ปัจจุบันคุณ{n}")
    age  = _val_or_ph(d.get("age", ""),              f"อายุคุณ{n}")

    sp_nick = _clean(d.get("คู่สมรส", d.get("spouse_nickname", "คู่สมรส"))) or "คู่สมรส"
    fs_sp   = _val_or_ph(d.get("fullname_spouse", ""), f"ชื่อ-นามสกุลคุณ{sp_nick}")

    ex_nick = _clean(d.get("ผู้จัดการมรดก", d.get("estate_executor", ""))) or "ผู้จัดการมรดก"
    fs_ex   = _val_or_ph(d.get("fullname_executor", ""),        f"ชื่อ-นามสกุล{ex_nick}")
    id_ex   = _val_or_ph(d.get("id_executor", ""),              f"เลขบัตรประชาชน{ex_nick}")
    fs_exb  = _val_or_ph(d.get("fullname_executor_backup", ""), "ชื่อ-นามสกุลผู้จัดการมรดกสำรอง")
    id_exb  = _val_or_ph(d.get("id_executor_backup", ""),       "เลขบัตรประชาชนผู้จัดการมรดกสำรอง")

    gp_nick  = _clean(d.get("Guardian หลัก",       d.get("guardian_primary", "")))       or "ผู้ปกครองหลัก"
    gb_nick  = _clean(d.get("Guardian สำรอง",       d.get("guardian_backup", "")))        or "ผู้ปกครองสำรอง"
    mgp_nick = _clean(d.get("Money Guardian หลัก",  d.get("money_guardian_primary", ""))) or "ผู้ดูแลทรัพย์สินหลัก"
    mgb_nick = _clean(d.get("Money Guardian สำรอง", d.get("money_guardian_backup", "")))  or "ผู้ดูแลทรัพย์สินสำรอง"

    fs_gp  = _val_or_ph(d.get("fullname_guardian_primary", ""),       f"ชื่อ-นามสกุล{gp_nick}")
    id_gp  = _val_or_ph(d.get("id_guardian_primary", ""),             f"เลขบัตรประชาชน{gp_nick}")
    fs_gb  = _val_or_ph(d.get("fullname_guardian_backup", ""),        f"ชื่อ-นามสกุล{gb_nick}")
    id_gb  = _val_or_ph(d.get("id_guardian_backup", ""),              f"เลขบัตรประชาชน{gb_nick}")
    fs_mgp = _val_or_ph(d.get("fullname_money_guardian_primary", ""), f"ชื่อ-นามสกุล{mgp_nick}")
    id_mgp = _val_or_ph(d.get("id_money_guardian_primary", ""),       f"เลขบัตรประชาชน{mgp_nick}")
    fs_mgb = _val_or_ph(d.get("fullname_money_guardian_backup", ""),  f"ชื่อ-นามสกุล{mgb_nick}")
    id_mgb = _val_or_ph(d.get("id_money_guardian_backup", ""),        f"เลขบัตรประชาชน{mgb_nick}")

    # ── document ───────────────────────────────────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.984)
        sec.bottom_margin = Inches(0.787)
        sec.left_margin   = Inches(1.181)
        sec.right_margin  = Inches(0.787)

    _para(doc, "พินัยกรรม", size=SZ_T, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "แบบธรรมดา ตามประมวลกฎหมายแพ่งและพาณิชย์ มาตรา ๑๖๕๖",
          size=SZ-1, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "ฉบับร่าง — โปรดตรวจสอบและลบบรรทัดนี้ก่อนพิมพ์จริง",
          size=SZ-1, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    _para(doc, "ทำที่ [ที่อยู่จัดทำเอกสาร]", space_after=4)
    _para(doc, "วันที่ [วันที่ทำพินัยกรรม] เดือน [เดือนที่ทำพินัยกรรม] พ.ศ. [ปีที่ทำพินัยกรรม]",
          space_after=12)

    # ข้อ ๑
    _clause(doc, "๑", "ผู้ทำพินัยกรรม")
    _para(doc,
          f"ข้าพเจ้า {fs} อายุ {age} ปี สัญชาติไทย "
          f"ถือบัตรประจำตัวประชาชนเลขที่ {ids} "
          f"อยู่บ้านเลขที่ {addr} "
          "ขณะทำพินัยกรรมฉบับนี้มีสติสัมปชัญญะสมบูรณ์ "
          "ไม่ได้อยู่ภายใต้การข่มขู่หรืออิทธิพลใดๆ "
          "ขอแสดงเจตนาอันเป็นคำสั่งครั้งสุดท้ายไว้ดังนี้")

    # ข้อ ๒
    _clause(doc, "๒", "การแต่งตั้งผู้จัดการมรดก")
    _para(doc,
          f"ข้าพเจ้าขอแต่งตั้งให้ {fs_ex} "
          f"ถือบัตรประจำตัวประชาชนเลขที่ {id_ex} "
          f"ซึ่งเป็น [ความสัมพันธ์กับคุณ{n}] "
          "เป็นผู้จัดการมรดกเพียงผู้เดียว มีอำนาจเต็มในการดำเนินการทุกอย่าง"
          "เกี่ยวกับกองมรดก รวมถึงการรับ ส่งมอบ ขาย โอน หรือชำระหนี้ของกองมรดก")
    _para(doc,
          f"หากผู้จัดการมรดกข้างต้นไม่สามารถปฏิบัติหน้าที่ได้ไม่ว่าด้วยเหตุใด "
          f"ให้ {fs_exb} ถือบัตรประจำตัวประชาชนเลขที่ {id_exb} เป็นผู้จัดการมรดกแทน")

    # ข้อ ๓
    _clause(doc, "๓", "การแบ่งทรัพย์สิน")
    for item in clause3_items:
        _subhead(doc, f"{item['num']}  {item['title']}")
        _para(doc, item['text'])

    # ข้อ ๔
    _clause(doc, "๔", "แผนสำรอง")
    _para(doc,
          f"หากคู่สมรสของข้าพเจ้า {fs_sp} "
          "ถึงแก่ความตายก่อนข้าพเจ้า หรือถึงแก่ความตายพร้อมกันกับข้าพเจ้า "
          "ให้ทรัพย์สินทั้งหมดตกเป็นของบุตรทั้งหมดของข้าพเจ้าในสัดส่วนเท่าๆ กัน")

    # ข้อ ๕
    _clause(doc, "๕", "การแต่งตั้งผู้ปกครองบุตรผู้เยาว์")
    _para(doc,
          "หากคู่สมรสของข้าพเจ้าถึงแก่ความตายพร้อมกันกับข้าพเจ้า "
          "ข้าพเจ้าขอแต่งตั้งบุคคลดังต่อไปนี้เพื่อดูแลบุตรทั้งหมดของข้าพเจ้า")
    _subhead(doc, "ผู้ปกครองด้านบุคคล")
    _para(doc, "มีหน้าที่ดูแลการเลี้ยงดู การศึกษา และความเป็นอยู่ของบุตร", space_after=4)
    _para(doc, f"หลัก: {fs_gp} ถือบัตรประจำตัวประชาชนเลขที่ {id_gp} ซึ่งเป็น [ความสัมพันธ์กับบุตร]")
    _para(doc, f"สำรอง: {fs_gb} ถือบัตรประจำตัวประชาชนเลขที่ {id_gb}")
    _subhead(doc, "ผู้ดูแลทรัพย์สินของบุตร")
    _para(doc,
          "มีหน้าที่บริหารและควบคุมทรัพย์สินของบุตรให้เป็นไปตามเจตนารมณ์ของพินัยกรรมฉบับนี้",
          space_after=4)
    _para(doc, f"หลัก: {fs_mgp} ถือบัตรประจำตัวประชาชนเลขที่ {id_mgp} ซึ่งเป็น [ความสัมพันธ์กับคุณ{n}]")
    _para(doc, f"สำรอง: {fs_mgb} ถือบัตรประจำตัวประชาชนเลขที่ {id_mgb}")
    _para(doc,
          "ข้าพเจ้าขอกำหนดเงื่อนไขว่า ผู้ปกครองด้านบุคคลและผู้ดูแลทรัพย์สิน"
          "ต้องเป็นคนละคนกันเสมอ หากบุคคลใดรับหน้าที่ทั้งสองบทบาทพร้อมกัน "
          "ให้ผู้จัดการมรดกยื่นคำร้องต่อศาลเพื่อแต่งตั้งบุคคลอื่นมาดำรงตำแหน่งที่ว่างแทน")

    # ข้อ ๖
    funeral = d.get("ความปรารถนางานศพ", d.get("funeral_wishes", "[รูปแบบพิธีและงบประมาณ]"))
    _clause(doc, "๖", "ความปรารถนาเกี่ยวกับพิธีศพ")
    _para(doc, f"ข้าพเจ้าประสงค์ให้จัด{funeral}")
    _para(doc,
          "ข้าพเจ้าขอแต่งตั้งให้ [ชื่อผู้รับผิดชอบจัดงานศพ] "
          "เป็นผู้รับผิดชอบจัดการพิธีศพ และมีอำนาจเบิกเงินจากกองมรดก"
          "เพื่อใช้จ่ายในการจัดพิธีได้ไม่เกิน [งบประมาณงานศพ] บาท")

    # ข้อ ๗
    _clause(doc, "๗", "เงื่อนไขทั่วไป")
    _para(doc,
          "พินัยกรรมฉบับนี้ให้ถือเป็นพินัยกรรมฉบับสุดท้ายของข้าพเจ้า "
          "และให้ยกเลิกพินัยกรรมฉบับอื่นๆ ทั้งหมดที่ได้ทำไว้ก่อนหน้านี้")

    # ลายเซ็น
    _spacer(doc, 14)
    _para(doc,
          "ข้าพเจ้าได้อ่านและเข้าใจข้อความในพินัยกรรมฉบับนี้โดยตลอดแล้ว "
          "จึงลงลายมือชื่อไว้เป็นสำคัญต่อหน้าพยาน")
    _spacer(doc, 6)
    _sig_line(doc, "ผู้ทำพินัยกรรม", f"[ชื่อ-นามสกุลคุณ{n}]")
    _spacer(doc, 10)
    _para(doc,
          "พยานรับรองว่าผู้ทำพินัยกรรมมีสติสัมปชัญญะสมบูรณ์ "
          "ได้ลงนามต่อหน้าพยานจริง และพยานทั้งสองมิใช่ผู้รับพินัยกรรม"
          "หรือคู่สมรสของผู้รับพินัยกรรม", space_after=4)
    _sig_line(doc, "พยานที่ ๑", "[ชื่อ-นามสกุลพยานที่ ๑]")
    _para(doc, "        บัตรประชาชนเลขที่ [เลขบัตรประชาชนพยานที่ ๑]  ที่อยู่ [ที่อยู่พยานที่ ๑]",
          size=SZ-1, space_before=0, space_after=4)
    _sig_line(doc, "พยานที่ ๒", "[ชื่อ-นามสกุลพยานที่ ๒]")
    _para(doc, "        บัตรประชาชนเลขที่ [เลขบัตรประชาชนพยานที่ ๒]  ที่อยู่ [ที่อยู่พยานที่ ๒]",
          size=SZ-1, space_before=0, space_after=16)
    _para(doc, "ผู้เขียน: พยัต จิรสุวรรณพงศ์",
          size=SZ-2, color=GRAY, space_before=0, space_after=0)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name); tmp.close()
    return tmp.name


# ── embed fonts ────────────────────────────────────────────────────────
def embed_fonts(docx_path: str) -> str:
    out = docx_path.replace(".docx", "_emb.docx")
    shutil.copy(docx_path, out)
    font_map = {
        "THSarabunNew":            os.path.join(FONT_DIR, "THSarabunNew.ttf"),
        "THSarabunNew-Bold":       os.path.join(FONT_DIR, "THSarabunNew Bold.ttf"),
        "THSarabunNew-Italic":     os.path.join(FONT_DIR, "THSarabunNew Italic.ttf"),
        "THSarabunNew-BoldItalic": os.path.join(FONT_DIR, "THSarabunNew BoldItalic.ttf"),
    }
    with zipfile.ZipFile(out, "a") as z:
        for name, path in font_map.items():
            arcname = f"word/fonts/{name}.ttf"
            try: z.getinfo(arcname)
            except KeyError:
                if os.path.exists(path): z.write(path, arcname)
    with zipfile.ZipFile(out, "r") as z:
        names = z.namelist()
        ft = z.read("word/fontTable.xml").decode()
        ct = z.read("[Content_Types].xml").decode()
        fr = (z.read("word/_rels/fontTable.xml.rels").decode()
              if "word/_rels/fontTable.xml.rels" in names else
              '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
              '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>')
        all_files = {n: z.read(n) for n in names}
    old = '<w:font w:name="TH Sarabun New"/>'
    new = ('<w:font w:name="TH Sarabun New">'
           '<w:embedRegular r:id="rIdF1"/><w:embedBold r:id="rIdF2"/>'
           '<w:embedItalic r:id="rIdF3"/><w:embedBoldItalic r:id="rIdF4"/></w:font>')
    ft = ft.replace(old, new) if old in ft else ft.replace("</w:fonts>", new + "</w:fonts>")
    rels = ''.join([
        '<Relationship Id="rIdF1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew.ttf"/>',
        '<Relationship Id="rIdF2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Bold.ttf"/>',
        '<Relationship Id="rIdF3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Italic.ttf"/>',
        '<Relationship Id="rIdF4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-BoldItalic.ttf"/>',
    ])
    fr = fr.replace("</Relationships>", rels + "</Relationships>")
    if "ttf" not in ct:
        ct = ct.replace("</Types>", '<Default Extension="ttf" ContentType="application/x-font-ttf"/></Types>')
    all_files["word/fontTable.xml"]            = ft.encode()
    all_files["word/_rels/fontTable.xml.rels"] = fr.encode()
    all_files["[Content_Types].xml"]            = ct.encode()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in all_files.items(): z.writestr(name, data)
    os.unlink(docx_path)
    return out


# ── entry point ────────────────────────────────────────────────────────
def build_will(client_data: dict) -> str:
    print("กำลัง generate ข้อ ๓...")
    clause3    = generate_clause3(client_data)
    print(f"ได้ {len(clause3)} sub-clause")
    raw_path   = build_will_docx(client_data, clause3)
    final_path = embed_fonts(raw_path)
    print("✅ สร้างพินัยกรรมเสร็จ")
    return final_path
