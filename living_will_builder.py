"""
living_will_builder.py — doc 4
หน้า 1: คู่มือ (Living Will คืออะไร + POA คืออะไร)
หน้า 2: หนังสือแสดงเจตนา Living Will (fill ชื่อจริง)
หน้า 3: หนังสือมอบอำนาจทั่วไป POA (fill ชื่อจริง)
"""
import os, re, json, tempfile, zipfile, shutil
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-5-20250929")

FONT     = "Cordia New"
BLACK    = RGBColor(0x00, 0x00, 0x00)
GRAY     = RGBColor(0x88, 0x88, 0x88)
RED      = RGBColor(0xAA, 0x00, 0x00)
NAVY     = RGBColor(0x1A, 0x3A, 0x5C)
GREEN    = RGBColor(0x1A, 0x7A, 0x4A)
LS       = 1.3
FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")

_NA = {'ยังไม่มี', '-', 'ยังไม่ได้กำหนด', 'ไม่ได้ระบุ', 'ไม่มี', ''}


# ── data helpers ──────────────────────────────────────────────────────
def _get(d, *keys):
    for k in keys:
        v = str(d.get(k, '') or '').strip()
        if v and v not in _NA:
            return v
    return ''

def _fill(d, placeholder, *keys):
    v = _get(d, *keys)
    return v if v else f"[{placeholder}]"

def _clean(s):
    return re.sub(r'\s*\(.*?\)', '', s or '').strip()


# ── text helpers ──────────────────────────────────────────────────────
def _run(para, text, bold=False, sz=13, color=None, underline=False):
    r = para.add_run(text)
    r.font.name = FONT; r.font.size = Pt(sz)
    r.bold = bold; r.font.color.rgb = color or BLACK
    if underline: r.underline = True
    return r

def _add_runs(para, text, sz=13):
    for part in re.split(r'(\[.*?\])', text):
        if not part: continue
        if part.startswith('['):
            _run(para, part, bold=True, sz=sz, color=RED, underline=True)
        else:
            _run(para, part, sz=sz)

def _p(doc, text, sz=13, bold=False, color=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    para.alignment = align
    if bold or color:
        _run(para, text, bold=bold, sz=sz, color=color)
    else:
        _add_runs(para, text, sz=sz)

def _head(doc, text, sz=15, sb=10, sa=3, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    _run(para, text, bold=True, sz=sz, color=color or BLACK)

def _bullet(doc, text, sz=13):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    para.paragraph_format.left_indent  = Pt(14)
    _run(para, "- ", sz=sz)
    _add_runs(para, text, sz=sz)

def _clause(doc, n, text, sz=13):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    _run(para, f"{n}.  ", bold=True, sz=sz)
    _add_runs(para, text, sz=sz)


# ── signature table ───────────────────────────────────────────────────
def _no_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0'); el.set(qn('w:color'),'auto')
        tcBdr.append(el)
    tcPr.append(tcBdr)

def _no_table_border(tbl):
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl_elem.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),"none"); el.set(qn("w:sz"),"0"); el.set(qn("w:color"),"auto")
        tblBdr.append(el)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None: tblPr.remove(old)
    tblPr.append(tblBdr)

def _sig_table(doc, pairs, sz=12, sb=8):
    n_cols   = len(pairs)
    col_w    = (9000 - 800 * (n_cols - 1)) // n_cols
    col_widths = []
    for i in range(n_cols):
        col_widths.append(col_w)
        if i < n_cols - 1: col_widths.append(800)

    tbl = doc.add_table(rows=3, cols=len(col_widths))
    tbl.style = "Table Grid"; _no_table_border(tbl)
    for i, w in enumerate(col_widths):
        for row in tbl.rows: row.cells[i].width = w * 635
    for row in tbl.rows:
        for cell in row.cells: _no_border(cell)

    tbl.rows[0].height = Pt(30)
    tbl.rows[1].height = Pt(2)
    tbl.rows[2].height = Pt(sz * LS + 2)
    content_cols = [i * 2 for i in range(n_cols)]

    for idx, (role, _) in enumerate(pairs):
        cp = tbl.rows[0].cells[content_cols[idx]].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cp.paragraph_format.space_before = Pt(sb if idx == 0 else 0)
        cp.paragraph_format.space_after  = Pt(0)
        cp.paragraph_format.line_spacing = Pt(sz * LS)
        _run(cp, role, sz=sz, color=GRAY)

    for idx in range(n_cols):
        cell = tbl.rows[1].cells[content_cols[idx]]
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcBdr = OxmlElement('w:tcBdr')
        for side in ['top','left','bottom','right']:
            el = OxmlElement(f'w:{side}')
            if side == 'top':
                el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'AAAAAA')
            else:
                el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0'); el.set(qn('w:color'),'auto')
            tcBdr.append(el)
        old = tcPr.find(qn('w:tcBdr'))
        if old is not None: tcPr.remove(old)
        tcPr.append(tcBdr)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    for idx, (_, name_ph) in enumerate(pairs):
        cp = tbl.rows[2].cells[content_cols[idx]].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(6)
        cp.paragraph_format.line_spacing = Pt(sz * LS)
        _run(cp, "(", sz=sz-1, color=GRAY)
        _add_runs(cp, name_ph, sz=sz-1)
        _run(cp, ")", sz=sz-1, color=GRAY)


# ── choose living will intent ─────────────────────────────────────────
def choose_intent(client_data: dict) -> str:
    lw = _get(client_data, "Living Will", "living_will")
    if not lw: return "B"
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    r = client.messages.create(
        model=MODEL, max_tokens=10,
        messages=[{"role": "user", "content":
            f'ตอบแค่ A หรือ B\nA=ไม่ยื้อชีวิต B=ยื้อชีวิตหรือไม่ชัดเจน\n"{lw}"\nตอบ:'}]
    )
    return "A" if r.content[0].text.strip().upper().startswith("A") else "B"


# ── PAGE 1: คู่มือ ────────────────────────────────────────────────────
def _build_page1(doc):
    _p(doc, "คู่มือเอกสารในชุดนี้",
       sz=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    _p(doc, "เอกสารชุดนี้ประกอบด้วย 2 ส่วน: (1) หนังสือแสดงเจตนาการรักษาพยาบาล และ (2) หนังสือมอบอำนาจ",
       sz=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=10)

    # Living Will
    _head(doc, "ส่วนที่ ๑ — หนังสือแสดงเจตนาการรักษาพยาบาล (Living Will)", sz=15, sb=4, sa=3, color=NAVY)

    _head(doc, "คืออะไร", sz=13, sb=4, sa=2)
    _p(doc,
       "เป็นเอกสารที่คุณบอกล่วงหน้าว่าต้องการหรือไม่ต้องการรับการรักษาแบบไหน "
       "ในกรณีที่ไม่สามารถตัดสินใจเองได้ เช่น หมดสติ หรืออยู่ในระยะสุดท้ายของชีวิต "
       "มีผลทางกฎหมายตาม มาตรา ๑๒ พ.ร.บ.สุขภาพแห่งชาติ พ.ศ. ๒๕๕๐", sa=3)

    _head(doc, "ข้อจำกัด", sz=13, sb=6, sa=2)
    _bullet(doc, "ใช้ได้เฉพาะกรณีอยู่ในระยะสุดท้าย หมดสติถาวร หรือไม่มีโอกาสฟื้นตัว เท่านั้น")
    _bullet(doc, "ไม่ครอบคลุมอุบัติเหตุทั่วไปที่ยังมีโอกาสฟื้นตัว")
    _bullet(doc, "แพทย์ต้องวินิจฉัยก่อนว่าอยู่ในสภาวะที่เอกสารนี้ใช้บังคับได้")

    _head(doc, "วิธีใช้", sz=13, sb=6, sa=2)
    _bullet(doc, "พกพาติดตัวหรือฝากไว้กับคนที่ไว้วางใจ และแจ้งให้ครอบครัวรู้ว่าอยู่ที่ไหน")
    _bullet(doc, "เมื่อเข้าโรงพยาบาล แจ้งแพทย์และยื่นหนังสือนี้ให้ทันที")
    _bullet(doc, "ผู้ตัดสินใจแทนที่ระบุไว้มีสิทธิ์ยื่นแทนได้ หากเจ้าของไม่สามารถยื่นเองได้")
    _bullet(doc, "ใช้ต้นฉบับเท่านั้น")

    _head(doc, "ลงทะเบียนออนไลน์ด้วย (แนะนำ)", sz=13, sb=6, sa=2)
    _p(doc,
       "ลงทะเบียนที่ e-livingwill.nationalhealth.or.th ฟรี "
       "เมื่อลงทะเบียนแล้ว แพทย์ทั่วประเทศสามารถตรวจสอบเจตนาของคุณได้ทันที "
       "แม้ไม่มีเอกสารฉบับกระดาษติดตัว ซึ่งสำคัญมากในกรณีฉุกเฉิน", sa=3)

    # POA
    _head(doc, "ส่วนที่ ๒ — หนังสือมอบอำนาจทั่วไป (POA)", sz=15, sb=10, sa=3, color=NAVY)

    _head(doc, "คืออะไร", sz=13, sb=4, sa=2)
    _p(doc,
       "เป็นเอกสารที่มอบสิทธิ์ให้คนอื่นทำธุรกรรมแทนในชีวิตประจำวัน "
       "ใช้เฉพาะกรณีที่คุณยังมีชีวิตอยู่แต่ทำเองไม่ได้ชั่วคราว เช่น ป่วย พักฟื้น หรือไม่สะดวกเดินทาง", sa=3)

    _head(doc, "ใช้ได้ในกรณีไหน", sz=13, sb=6, sa=2)
    _bullet(doc, "จ่ายค่าสาธารณูปโภค เช่น ค่าไฟ ค่าน้ำ ค่าโทรศัพท์")
    _bullet(doc, "ติดต่อหน่วยงานราชการแทน เช่น อำเภอ สรรพากร ประกันสังคม")
    _bullet(doc, "ลงนามในสัญญาทั่วไป เช่น สัญญาเช่า สัญญาซื้อขาย")
    _bullet(doc, "รับส่งเอกสารแทน และติดต่อบริษัทต่างๆ")

    _head(doc, "ข้อจำกัดสำคัญ", sz=13, sb=6, sa=2)
    _bullet(doc, "สิ้นผลทันทีเมื่อผู้มอบอำนาจเสียชีวิต — ต้องใช้พินัยกรรมแทน")
    _bullet(doc, "ใช้ไม่ได้กับกรมที่ดิน ต้องใช้แบบฟอร์ม ท.ด.21 โดยเฉพาะ")
    _bullet(doc, "ธนาคารบางแห่งมีแบบฟอร์มของตัวเอง อาจต้องกรอกเพิ่มเติม")
    _bullet(doc, "ใช้ต้นฉบับเท่านั้น พกบัตรประชาชนของผู้รับมอบอำนาจไปด้วยทุกครั้ง")

    _head(doc, "จำนวนชุดที่ควรเตรียม", sz=13, sb=6, sa=2)
    _p(doc,
       "ควรเตรียมอย่างน้อย 3-4 ชุด สำหรับธนาคาร บริษัทประกัน หน่วยงานราชการ "
       "และสำรองเก็บที่บ้าน 1 ชุด โดยต้องเซ็นชื่อต้นฉบับทุกฉบับพร้อมกันต่อหน้าพยานในคราวเดียว", sa=0)


# ── PAGE 2: Living Will ───────────────────────────────────────────────
def _build_page2(doc, client_data: dict, intent: str):
    d = client_data
    n  = _get(d, "ชื่อเล่น", "nickname") or "ผู้ทำหนังสือ"
    sp = _clean(_get(d, "คู่สมรส", "spouse_nickname")) or "คู่สมรส"

    fs   = _fill(d, f"ชื่อ-นามสกุลคุณ{n}",  "ชื่อ-นามสกุลจริงเจ้าของแผน", "fullname_self")
    ids  = _fill(d, f"เลขบัตรประชาชนคุณ{n}", "เลขบัตรประชาชนเจ้าของแผน", "id_self")
    addr = _fill(d, f"ที่อยู่ปัจจุบันคุณ{n}", "ที่อยู่ปัจจุบัน", "address_self")
    age  = _get(d, "อายุ", "age") or f"อายุคุณ{n}"
    fs_sp = _fill(d, f"ชื่อ-นามสกุลคุณ{sp}", "ชื่อ-นามสกุลจริงคู่สมรส", "fullname_spouse")

    _p(doc, "หนังสือแสดงเจตนาเกี่ยวกับการรักษาพยาบาล",
       sz=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    _p(doc, "ตามมาตรา ๑๒ แห่งพระราชบัญญัติสุขภาพแห่งชาติ พ.ศ. ๒๕๕๐",
       sz=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    _p(doc, "ฉบับร่าง — โปรดตรวจสอบและลบบรรทัดนี้ก่อนพิมพ์จริง",
       sz=11, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, sa=8)

    _p(doc, "ทำที่ [ที่อยู่จัดทำเอกสาร]", sa=1)
    _p(doc, "วันที่ [วันที่ทำหนังสือ] เดือน [เดือน] พ.ศ. [ปี]", sa=6)

    _p(doc,
       f"ข้าพเจ้า {fs} อายุ {age} ปี สัญชาติไทย "
       f"ถือบัตรประจำตัวประชาชนเลขที่ {ids} "
       f"อยู่บ้านเลขที่ {addr}", sa=4)
    _p(doc,
       "ขณะที่ข้าพเจ้ามีสติสัมปชัญญะสมบูรณ์ดี ข้าพเจ้าขอแสดงเจตนาไว้ล่วงหน้าว่า "
       "หากข้าพเจ้าอยู่ในสภาวะใดสภาวะหนึ่งดังต่อไปนี้", sa=4)
    _bullet(doc, "อยู่ในระยะสุดท้ายของชีวิต ซึ่งแพทย์วินิจฉัยว่าไม่สามารถรักษาให้หายได้")
    _bullet(doc, "อยู่ในภาวะหมดสติถาวรซึ่งไม่มีโอกาสฟื้นคืนสติ")
    _bullet(doc, "อยู่ในภาวะที่การรักษาไม่มีโอกาสทำให้กลับมาใช้ชีวิตได้อย่างมีคุณภาพ")

    _head(doc, "เจตนาของข้าพเจ้า", sz=14, sb=6, sa=3)
    if intent == "A":
        _p(doc, "ข้าพเจ้าไม่ประสงค์จะรับการรักษาที่มีเป้าหมายเพื่อยืดการตายเท่านั้น โดยเฉพาะ", sa=4)
        _bullet(doc, "การใส่เครื่องช่วยหายใจในกรณีที่ไม่มีโอกาสฟื้นตัว")
        _bullet(doc, "การกระตุ้นหัวใจในกรณีที่แพทย์วินิจฉัยว่าไม่มีโอกาสฟื้นตัว")
        _bullet(doc, "การให้อาหารทางสายยางในกรณีที่อยู่ในภาวะหมดสติถาวร")
        _bullet(doc, "การรักษาที่ก่อให้เกิดความทรมานโดยไม่มีประโยชน์ต่อการฟื้นตัว")
        _p(doc,
           "ข้าพเจ้าประสงค์จะได้รับการดูแลแบบประคับประคองเพื่อบรรเทาความเจ็บปวด "
           "และให้สามารถจากไปได้อย่างสงบ รวมถึงยาบรรเทาปวดในปริมาณที่เพียงพอ",
           sb=4, sa=4)
    else:
        _p(doc,
           "ข้าพเจ้าประสงค์จะรับการรักษาพยาบาลอย่างเต็มที่ทุกวิธีที่แพทย์เห็นสมควร "
           "รวมถึงการใช้เครื่องช่วยหายใจ การกระตุ้นหัวใจ และการรักษาอื่นๆ "
           "เพื่อยืดชีวิตให้นานที่สุดเท่าที่จะเป็นไปได้", sa=4)

    _head(doc, "ผู้ตัดสินใจแทน", sz=14, sb=6, sa=2)
    _p(doc, "หากข้าพเจ้าไม่สามารถแสดงเจตนาด้วยตนเองได้ ขอมอบหมายให้บุคคลต่อไปนี้เป็นผู้ตัดสินใจแทน", sa=4)
    _p(doc, f"หลัก: {fs_sp}  ความสัมพันธ์: [ความสัมพันธ์กับคุณ{n}]  โทร: [เบอร์โทรคุณ{sp}]", sa=2)
    _p(doc, "สำรอง: [ชื่อ-นามสกุลผู้ตัดสินใจแทนสำรอง]  ความสัมพันธ์: [ความสัมพันธ์]  โทร: [เบอร์โทร]", sa=6)

    _p(doc,
       "ข้าพเจ้าได้ทำหนังสือฉบับนี้ด้วยความสมัครใจ มีสติสัมปชัญญะสมบูรณ์ "
       "และขอให้แพทย์และบุคลากรทางการแพทย์ทุกท่านปฏิบัติตามเจตนาที่ระบุไว้", sa=4)

    _sig_table(doc, [("ผู้แสดงเจตนา", f"[ชื่อ-นามสกุลคุณ{n}]")], sz=12, sb=4)
    _p(doc, "พยานรับรองว่าผู้แสดงเจตนามีสติสัมปชัญญะสมบูรณ์และลงนามด้วยความสมัครใจ", sb=8, sa=2)
    _sig_table(doc, [
        ("พยานที่ ๑", "[ชื่อ-นามสกุลพยานที่ ๑]"),
        ("พยานที่ ๒", "[ชื่อ-นามสกุลพยานที่ ๒]"),
    ], sz=12, sb=4)

    _p(doc,
       "💡 แนะนำ: ลงทะเบียนออนไลน์เพิ่มเติมที่ e-livingwill.nationalhealth.or.th "
       "เพื่อให้แพทย์ทั่วประเทศเข้าถึงเจตนาของคุณได้ทันทีในกรณีฉุกเฉิน โดยไม่ต้องรอเอกสารกระดาษ",
       sz=11, color=GREEN, sb=8, sa=0)


# ── PAGE 3: POA ───────────────────────────────────────────────────────
def _build_page3(doc, client_data: dict):
    d = client_data
    n  = _get(d, "ชื่อเล่น", "nickname") or "ผู้มอบอำนาจ"
    sp = _clean(_get(d, "คู่สมรส", "spouse_nickname")) or "คู่สมรส"

    fs_self   = _fill(d, f"ชื่อ-นามสกุลคุณ{n}",  "ชื่อ-นามสกุลจริงเจ้าของแผน", "fullname_self")
    id_self   = _fill(d, f"เลขบัตรประชาชนคุณ{n}", "เลขบัตรประชาชนเจ้าของแผน", "id_self")
    addr_self = _fill(d, f"ที่อยู่ปัจจุบันคุณ{n}", "ที่อยู่ปัจจุบัน", "address_self")
    age_self  = _get(d, "อายุ", "age") or f"อายุคุณ{n}"
    fs_sp     = _fill(d, "ชื่อ-นามสกุลคู่สมรส",   "ชื่อ-นามสกุลจริงคู่สมรส", "fullname_spouse")
    id_sp     = _fill(d, "เลขบัตรประชาชนคู่สมรส", "เลขบัตรประชาชนคู่สมรส", "id_spouse")
    addr_sp   = _fill(d, "ที่อยู่ปัจจุบันคู่สมรส", "ที่อยู่ปัจจุบันคู่สมรส", "address_spouse")
    age_sp    = _get(d, "อายุคู่สมรส", "spouse_age") or f"อายุคุณ{sp}"

    _p(doc, "หนังสือมอบอำนาจ",
       sz=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    _p(doc, "มอบอำนาจทั่วไปในชีวิตประจำวัน",
       sz=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    _p(doc, "ฉบับร่าง — โปรดตรวจสอบและลบบรรทัดนี้ก่อนพิมพ์จริง",
       sz=11, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, sa=8)

    _p(doc, "ทำที่ [ที่อยู่จัดทำเอกสาร]", sa=1)
    _p(doc, "วันที่ [วันที่ทำหนังสือ] เดือน [เดือน] พ.ศ. [ปี]", sa=6)

    _head(doc, "ผู้มอบอำนาจ", sz=14, sb=4, sa=2)
    _p(doc,
       f"ข้าพเจ้า {fs_self} อายุ {age_self} ปี สัญชาติไทย "
       f"ถือบัตรประจำตัวประชาชนเลขที่ {id_self} "
       f"อยู่บ้านเลขที่ {addr_self} "
       "ซึ่งต่อไปในหนังสือนี้จะเรียกว่า ผู้มอบอำนาจ")

    _head(doc, "ผู้รับมอบอำนาจ", sz=14, sb=4, sa=2)
    _p(doc,
       f"{fs_sp} อายุ {age_sp} ปี สัญชาติไทย "
       f"ถือบัตรประจำตัวประชาชนเลขที่ {id_sp} "
       f"อยู่บ้านเลขที่ {addr_sp} "
       "ซึ่งต่อไปในหนังสือนี้จะเรียกว่า ผู้รับมอบอำนาจ")

    _head(doc, "ขอบเขตการมอบอำนาจ", sz=14, sb=4, sa=2)
    _p(doc, "ผู้มอบอำนาจขอมอบอำนาจให้ผู้รับมอบอำนาจกระทำการแทนในกิจการต่อไปนี้", sa=2)

    for num, text in zip(["๑","๒","๓","๔","๕"], [
        "ชำระค่าสาธารณูปโภค เช่น ค่าไฟฟ้า ค่าน้ำประปา ค่าโทรศัพท์ ค่าอินเทอร์เน็ต",
        "ติดต่อหน่วยงานราชการแทน เช่น สำนักงานเขต อำเภอ กรมสรรพากร สำนักงานประกันสังคม",
        "ลงนามในสัญญาทั่วไป เช่น สัญญาเช่า สัญญาซื้อขายทรัพย์สินมูลค่าไม่เกิน [จำนวนเงิน] บาท",
        "รับส่งเอกสารและพัสดุแทนผู้มอบอำนาจ",
        "ติดต่อและประสานงานกับบริษัทและองค์กรต่างๆ ในนามของผู้มอบอำนาจ",
    ]):
        _clause(doc, num, text)

    _head(doc, "เงื่อนไขสำคัญ", sz=14, sb=6, sa=2)
    _p(doc,
       "หนังสือฉบับนี้มีผลเฉพาะกรณีที่ผู้มอบอำนาจยังมีชีวิตอยู่แต่ไม่สามารถดำเนินการเองได้ชั่วคราว "
       "ผู้มอบอำนาจสามารถเพิกถอนได้ทุกเมื่อโดยแจ้งเป็นลายลักษณ์อักษร",
       sa=2)
    _p(doc,
       "⚠️ หนังสือมอบอำนาจฉบับนี้สิ้นผลทันทีเมื่อผู้มอบอำนาจเสียชีวิต "
       "สำหรับการจัดการมรดก ให้ใช้พินัยกรรมและติดต่อผู้จัดการมรดกแทน",
       sz=12, bold=True, color=RED, sa=6)

    _p(doc, "เพื่อเป็นหลักฐาน คู่สัญญาทั้งสองฝ่ายได้ลงลายมือชื่อไว้ต่อหน้าพยาน", sa=2)
    _sig_table(doc, [
        ("ผู้มอบอำนาจ",    f"[ชื่อ-นามสกุลคุณ{n}]"),
        ("ผู้รับมอบอำนาจ", f"[ชื่อ-นามสกุลคู่สมรส]"),
    ], sz=12, sb=6)
    _p(doc, "พยานรับรองว่าคู่สัญญาทั้งสองฝ่ายได้ลงนามต่อหน้าพยานจริง", sb=8, sa=2)
    _sig_table(doc, [
        ("พยานที่ ๑", "[ชื่อ-นามสกุลพยานที่ ๑]"),
        ("พยานที่ ๒", "[ชื่อ-นามสกุลพยานที่ ๒]"),
    ], sz=12, sb=4)


# ── build full docx ───────────────────────────────────────────────────
def build_lw_docx(client_data: dict, intent: str) -> str:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.787)
        section.bottom_margin = Inches(0.787)
        section.left_margin   = Inches(1.181)
        section.right_margin  = Inches(0.787)

    _build_page1(doc)
    doc.add_page_break()
    _build_page2(doc, client_data, intent)
    doc.add_page_break()
    _build_page3(doc, client_data)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name); tmp.close()
    return tmp.name


# ── embed fonts ───────────────────────────────────────────────────────
def embed_fonts(docx_path: str) -> str:
    out = docx_path.replace(".docx", "_emb.docx")
    shutil.copy(docx_path, out)
    font_map = {
        "THSarabunNew":            os.path.join(FONT_DIR, "THSarabunNew.ttf"),
        "THSarabunNew-Bold":       os.path.join(FONT_DIR, "THSarabunNew Bold.ttf"),
        "THSarabunNew-Italic":     os.path.join(FONT_DIR, "THSarabunNew Italic.ttf"),
        "THSarabunNew-BoldItalic": os.path.join(FONT_DIR, "THSarabunNew BoldItalic.ttf"),
    }
    with zipfile.ZipFile(out, "a") as z:
        for name, path in font_map.items():
            arcname = f"word/fonts/{name}.ttf"
            try: z.getinfo(arcname)
            except KeyError:
                if os.path.exists(path): z.write(path, arcname)
    with zipfile.ZipFile(out, "r") as z:
        names = z.namelist()
        ft = z.read("word/fontTable.xml").decode()
        ct = z.read("[Content_Types].xml").decode()
        fr = (z.read("word/_rels/fontTable.xml.rels").decode()
              if "word/_rels/fontTable.xml.rels" in names else
              '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
              '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>')
        all_files = {n: z.read(n) for n in names}
    old = '<w:font w:name="TH Sarabun New"/>'
    new = ('<w:font w:name="TH Sarabun New">'
           '<w:embedRegular r:id="rIdF1"/><w:embedBold r:id="rIdF2"/>'
           '<w:embedItalic r:id="rIdF3"/><w:embedBoldItalic r:id="rIdF4"/></w:font>')
    ft = ft.replace(old, new) if old in ft else ft.replace("</w:fonts>", new + "</w:fonts>")
    rels = ''.join([
        '<Relationship Id="rIdF1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew.ttf"/>',
        '<Relationship Id="rIdF2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Bold.ttf"/>',
        '<Relationship Id="rIdF3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Italic.ttf"/>',
        '<Relationship Id="rIdF4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-BoldItalic.ttf"/>',
    ])
    fr = fr.replace("</Relationships>", rels + "</Relationships>")
    if "ttf" not in ct:
        ct = ct.replace("</Types>", '<Default Extension="ttf" ContentType="application/x-font-ttf"/></Types>')
    all_files["word/fontTable.xml"]            = ft.encode()
    all_files["word/_rels/fontTable.xml.rels"] = fr.encode()
    all_files["[Content_Types].xml"]            = ct.encode()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in all_files.items(): z.writestr(name, data)
    os.unlink(docx_path)
    return out


# ── entry point ───────────────────────────────────────────────────────
def build_living_will(client_data: dict) -> str:
    print("กำลังเลือก template Living Will...")
    intent = choose_intent(client_data)
    print(f"✅ เลือกแบบ {'ก (ไม่ยื้อ)' if intent == 'A' else 'ข (ยื้อ)'}")
    raw_path   = build_lw_docx(client_data, intent)
    final_path = embed_fonts(raw_path)
    print("✅ สร้างหนังสือแสดงเจตนา + POA เสร็จ")
    return final_path
