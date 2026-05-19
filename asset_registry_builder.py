"""
asset_registry_builder.py — doc 3
บัญชีทรัพย์สินและคู่มือเข้าถึง
- หมวด 1: ทรัพย์สินและหนี้สิน (gen จาก JSON)
- หมวด 2: ประกันทุกกรมธรรม์ (gen จาก JSON)
- หมวด 3: บัญชีธนาคาร / app / รหัส / crypto (ตารางเปล่ากรอกมือ)
- หมวด 4: วิธีจัดการ (fixed template)
"""
import os, re, tempfile, zipfile, shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT     = "Cordia New"
BLACK    = RGBColor(0x00, 0x00, 0x00)
GRAY     = RGBColor(0x88, 0x88, 0x88)
RED      = RGBColor(0xAA, 0x00, 0x00)
NAVY     = RGBColor(0x1A, 0x3A, 0x5C)
AMBER    = RGBColor(0x7A, 0x58, 0x00)
AMBER_BG = "FFFBEA"
LS       = 1.3
FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")

_NA = {'ยังไม่มี', '-', 'ยังไม่ได้กำหนด', 'ไม่ได้ระบุ', 'ไม่มี', ''}


# ── data helpers ──────────────────────────────────────────────────────
def _get(d, *keys):
    for k in keys:
        v = str(d.get(k, '') or '').strip()
        if v and v not in _NA:
            return v
    return ''

def _has(d, *keys):
    return bool(_get(d, *keys))


# ── text helpers ──────────────────────────────────────────────────────
def _run(para, text, bold=False, sz=13, color=None, underline=False):
    r = para.add_run(text)
    r.font.name = FONT; r.font.size = Pt(sz)
    r.bold = bold; r.font.color.rgb = color or BLACK
    if underline: r.underline = True
    return r

def _p(doc, text, sz=13, bold=False, color=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    para.alignment = align
    _run(para, text, bold=bold, sz=sz, color=color)
    return para

def _section(doc, n, title, sz=15, sb=14, sa=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    _run(para, f"หมวด {n} — {title}", bold=True, sz=sz, color=NAVY)

def _row(doc, label, value, sz=12):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    para.paragraph_format.left_indent  = Pt(12)
    _run(para, f"{label}: ", bold=True, sz=sz)
    _run(para, value or "—", sz=sz, color=GRAY if not value else BLACK)

def _note_box(doc, text, sz=11):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), AMBER_BG); shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
    tcBdr = OxmlElement('w:tcBdr')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:color'),'D4A000'); tcBdr.append(el)
    tcPr.append(tcBdr)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after  = Pt(4)
    cp.paragraph_format.line_spacing = Pt(sz * LS)
    _run(cp, text, sz=sz, color=AMBER)


# ── table helpers ─────────────────────────────────────────────────────
def _tbl_header(tbl, headers, sz=11):
    row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(h); r.font.name = FONT
        r.font.size = Pt(sz); r.bold = True
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1E3A5F"); shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

def _tbl_row(tbl, row_idx, cells, sz=11):
    row = tbl.rows[row_idx]
    bg = "F5F7FA" if row_idx % 2 == 1 else "FFFFFF"
    for i, txt in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = Pt(sz * LS)
        r = p.add_run(txt); r.font.name = FONT
        r.font.size = Pt(sz)
        r.font.color.rgb = GRAY if txt.startswith("____") else BLACK
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg); shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

def _blank_table(doc, headers, n_rows, sz=11):
    """ตารางเปล่า n_rows แถว"""
    tbl = doc.add_table(rows=n_rows+1, cols=len(headers))
    tbl.style = "Table Grid"
    _tbl_header(tbl, headers, sz=sz)
    blank = ["________________________"] * len(headers)
    for i in range(n_rows):
        _tbl_row(tbl, i+1, blank, sz=sz)


# ── build docx ────────────────────────────────────────────────────────
def build_asset_registry_docx(client_data: dict) -> str:
    d  = client_data
    n  = _get(d, "ชื่อเล่น", "nickname") or "ผู้ทำเอกสาร"

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.787)
        section.bottom_margin = Inches(0.787)
        section.left_margin   = Inches(1.181)
        section.right_margin  = Inches(0.787)

    # ── title ─────────────────────────────────────────────────────────
    _p(doc, f"บัญชีทรัพย์สินและคู่มือเข้าถึงของคุณ{n}",
       sz=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    _p(doc, "เอกสารนี้มีข้อมูลความลับสูง — กรอกข้อมูลจริงและเก็บในที่ปลอดภัยก่อนใช้",
       sz=11, color=RED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)

    _note_box(doc,
        "⚠️  เอกสารฉบับนี้มีข้อมูลความลับสูง ควรเก็บแยกจากเอกสารอื่นในที่ปลอดภัย "
        "และแจ้งให้รู้เฉพาะผู้จัดการมรดกที่ไว้วางใจได้เท่านั้น")

    # ── หมวด 1: ทรัพย์สินและหนี้สิน ─────────────────────────────────
    _section(doc, "๑", "ทรัพย์สินและหนี้สิน")

    # เงินสด
    cash = _get(d, "เงินสด", "assets_cash")
    if cash:
        _row(doc, "เงินสดและเงินฝาก", cash)

    # อสังหาริมทรัพย์
    prop = _get(d, "อสังหาริมทรัพย์", "assets_property")
    if prop:
        _row(doc, "อสังหาริมทรัพย์", prop)

    # การลงทุน
    inv = _get(d, "การลงทุน", "assets_investment")
    if inv:
        _row(doc, "หุ้น / กองทุน / การลงทุน", inv)

    # คริปโต
    crypto = _get(d, "คริปโต", "assets_crypto_wallet")
    if crypto:
        _row(doc, "สินทรัพย์ดิจิทัล", crypto)

    # ประกันสะสม
    sav = _get(d, "ประกันสะสม", "assets_insurance_savings")
    if sav:
        _row(doc, "ประกันสะสมทรัพย์", sav)

    # กิจการ
    biz = _get(d, "กิจการ", "assets_business")
    if biz:
        _row(doc, "กิจการ / ธุรกิจ", biz)

    # ของมีค่า
    val = _get(d, "ของมีค่า", "assets_valuables")
    if val:
        _row(doc, "ทองคำ / ของมีค่า", val)

    # หนี้สิน
    debt = _get(d, "หนี้สิน", "debt")
    _row(doc, "หนี้สิน", debt if debt else "ไม่มีหนี้สิน")

    # ค้ำประกัน
    guarantor = _get(d, "ค้ำประกัน", "guarantor")
    if guarantor:
        _row(doc, "ค้ำประกัน", guarantor)

    # ── หมวด 2: ประกันทุกกรมธรรม์ ───────────────────────────────────
    _section(doc, "๒", "ประกันทุกกรมธรรม์")

    life = _get(d, "ประกันชีวิต", "insurance_life")
    if life:
        _row(doc, "ประกันชีวิต", life)

    health = _get(d, "ประกันสุขภาพ", "insurance_health")
    if health:
        _row(doc, "ประกันสุขภาพ", health)

    group = _get(d, "ประกันกลุ่ม", "insurance_group")
    if group:
        _row(doc, "ประกันกลุ่ม", group)

    welfare = _get(d, "สวัสดิการ", "welfare")
    if welfare:
        _row(doc, "สวัสดิการอื่นๆ", welfare)

    _p(doc, "กรมธรรม์แต่ละฉบับ — กรอกเพิ่มเติมด้วยมือ",
       sz=11, color=GRAY, sb=6, sa=2)

    tbl2 = doc.add_table(rows=5, cols=5); tbl2.style = "Table Grid"
    _tbl_header(tbl2, ["บริษัทประกัน", "เลขกรมธรรม์", "ทุนประกัน (บาท)", "ผู้รับประโยชน์", "เบอร์ติดต่อ"])
    for i in range(1, 5):
        _tbl_row(tbl2, i, ["________________________"]*5)

    # ── หมวด 3: บัญชีธนาคารและการเข้าถึง ────────────────────────────
    _section(doc, "๓", "บัญชีธนาคารและการเข้าถึง")
    _note_box(doc,
        "⚠️  กรอกด้วยลายมือบนกระดาษ ใส่ซองปิดผนึก เก็บในเซฟ "
        "ห้ามพิมพ์ ถ่ายรูป หรือส่งทางออนไลน์เด็ดขาด")

    _p(doc, "บัญชีธนาคาร", sz=13, bold=True, sb=8, sa=2)
    tbl3 = doc.add_table(rows=5, cols=4); tbl3.style = "Table Grid"
    _tbl_header(tbl3, ["ธนาคาร", "เลขบัญชี", "ประเภทบัญชี", "รหัส ATM"])
    for i in range(1, 5):
        _tbl_row(tbl3, i, ["________________________"]*4)

    _p(doc, "Mobile Banking / App ธนาคาร", sz=13, bold=True, sb=10, sa=2)
    tbl4 = doc.add_table(rows=4, cols=4); tbl4.style = "Table Grid"
    _tbl_header(tbl4, ["ธนาคาร", "Username / เบอร์ลงทะเบียน", "PIN app", "วิธี reset"])
    for i in range(1, 4):
        _tbl_row(tbl4, i, ["________________________"]*4)

    _p(doc, "อีเมลหลัก", sz=13, bold=True, sb=10, sa=2)
    tbl5 = doc.add_table(rows=3, cols=3); tbl5.style = "Table Grid"
    _tbl_header(tbl5, ["Email address", "Password", "วิธี 2FA / Recovery"])
    for i in range(1, 3):
        _tbl_row(tbl5, i, ["________________________"]*3)

    _p(doc, "Platform / บัญชีสำคัญ", sz=13, bold=True, sb=10, sa=2)
    tbl6 = doc.add_table(rows=5, cols=4); tbl6.style = "Table Grid"
    _tbl_header(tbl6, ["Platform", "Username / Email", "Password", "วิธีเข้าถึง 2FA"])
    for i in range(1, 5):
        _tbl_row(tbl6, i, ["________________________"]*4)

    _p(doc, "Crypto Wallet (ถ้ามี)", sz=13, bold=True, sb=10, sa=2)
    tbl7 = doc.add_table(rows=3, cols=3); tbl7.style = "Table Grid"
    _tbl_header(tbl7, ["Wallet / Exchange", "Seed phrase เก็บที่", "วิธีเข้าถึง"])
    for i in range(1, 3):
        _tbl_row(tbl7, i, ["________________________"]*3)

    _p(doc, "รหัสอุปกรณ์", sz=13, bold=True, sb=10, sa=2)
    tbl8 = doc.add_table(rows=4, cols=3); tbl8.style = "Table Grid"
    _tbl_header(tbl8, ["อุปกรณ์", "PIN / Password", "Backup / Recovery"])
    for i, label in enumerate(["โทรศัพท์มือถือ", "คอมพิวเตอร์", "Password Manager"], 1):
        _tbl_row(tbl8, i, [label, "________________________", "________________________"])

    # ── หมวด 4: วิธีจัดการ ───────────────────────────────────────────
    _section(doc, "๔", "วิธีจัดการทรัพย์สินแต่ละประเภท")

    items = [
        ("เงินฝากธนาคาร",
         "แสดงใบมรณบัตร + คำสั่งศาลแต่งตั้งผู้จัดการมรดก ติดต่อที่สาขาธนาคาร "
         "หรือโทรสายด่วนธนาคารเพื่อแจ้งอายัดบัญชีก่อน"),
        ("หุ้น / กองทุนรวม",
         "ติดต่อโบรกเกอร์หรือบริษัทจัดการกองทุน แสดงใบมรณบัตร + คำสั่งศาล "
         "ขอโอนหน่วยลงทุนหรือถอนขายตามคำสั่งพินัยกรรม"),
        ("ประกันชีวิต",
         "ติดต่อบริษัทประกันโดยตรง ไม่ต้องรอคำสั่งศาล "
         "แสดงใบมรณบัตร + กรมธรรม์ + บัตรประชาชนผู้รับประโยชน์ "
         "ประกันชีวิตจ่ายตรงให้ผู้รับประโยชน์ภายใน 15-30 วันทำการ"),
        ("อสังหาริมทรัพย์",
         "ต้องมีคำสั่งศาลแต่งตั้งผู้จัดการมรดกก่อน จึงจะโอนกรรมสิทธิ์ได้ "
         "ติดต่อกรมที่ดิน โทร 02-141-5555 เตรียมโฉนด + คำสั่งศาล + บัตรประชาชน"),
        ("กิจการ / บริษัท",
         "ติดต่อกรมพัฒนาธุรกิจการค้า แสดงใบมรณบัตร + คำสั่งศาล "
         "เพื่อจดทะเบียนเปลี่ยนแปลงกรรมการหรือผู้ถือหุ้น"),
        ("สินทรัพย์ดิจิทัล",
         "ต้องมี Seed phrase หรือ Private key ในการเข้าถึง "
         "ห้ามเปิดเผยกับใครทางออนไลน์ ปรึกษาผู้เชี่ยวชาญก่อนดำเนินการ"),
    ]

    for label, desc in items:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after  = Pt(2)
        para.paragraph_format.line_spacing = Pt(13 * LS)
        para.paragraph_format.left_indent  = Pt(12)
        _run(para, f"{label}: ", bold=True, sz=12)
        _run(para, desc, sz=12)

    _p(doc, "ผู้เขียน: พยัต จิรสุวรรณพงศ์",
       sz=11, color=GRAY, sb=20, sa=0)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name); tmp.close()
    return tmp.name


# ── embed fonts ───────────────────────────────────────────────────────
def embed_fonts(docx_path: str) -> str:
    out = docx_path.replace(".docx", "_emb.docx")
    shutil.copy(docx_path, out)
    font_map = {
        "THSarabunNew":            os.path.join(FONT_DIR, "THSarabunNew.ttf"),
        "THSarabunNew-Bold":       os.path.join(FONT_DIR, "THSarabunNew Bold.ttf"),
        "THSarabunNew-Italic":     os.path.join(FONT_DIR, "THSarabunNew Italic.ttf"),
        "THSarabunNew-BoldItalic": os.path.join(FONT_DIR, "THSarabunNew BoldItalic.ttf"),
    }
    with zipfile.ZipFile(out, "a") as z:
        for name, path in font_map.items():
            arcname = f"word/fonts/{name}.ttf"
            try: z.getinfo(arcname)
            except KeyError:
                if os.path.exists(path): z.write(path, arcname)
    with zipfile.ZipFile(out, "r") as z:
        names = z.namelist()
        ft = z.read("word/fontTable.xml").decode()
        ct = z.read("[Content_Types].xml").decode()
        fr = (z.read("word/_rels/fontTable.xml.rels").decode()
              if "word/_rels/fontTable.xml.rels" in names else
              '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
              '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>')
        all_files = {n: z.read(n) for n in names}
    old = '<w:font w:name="TH Sarabun New"/>'
    new = ('<w:font w:name="TH Sarabun New">'
           '<w:embedRegular r:id="rIdF1"/><w:embedBold r:id="rIdF2"/>'
           '<w:embedItalic r:id="rIdF3"/><w:embedBoldItalic r:id="rIdF4"/></w:font>')
    ft = ft.replace(old, new) if old in ft else ft.replace("</w:fonts>", new + "</w:fonts>")
    rels = ''.join([
        '<Relationship Id="rIdF1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew.ttf"/>',
        '<Relationship Id="rIdF2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Bold.ttf"/>',
        '<Relationship Id="rIdF3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Italic.ttf"/>',
        '<Relationship Id="rIdF4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-BoldItalic.ttf"/>',
    ])
    fr = fr.replace("</Relationships>", rels + "</Relationships>")
    if "ttf" not in ct:
        ct = ct.replace("</Types>", '<Default Extension="ttf" ContentType="application/x-font-ttf"/></Types>')
    all_files["word/fontTable.xml"]            = ft.encode()
    all_files["word/_rels/fontTable.xml.rels"] = fr.encode()
    all_files["[Content_Types].xml"]            = ct.encode()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in all_files.items(): z.writestr(name, data)
    os.unlink(docx_path)
    return out


# ── entry point ───────────────────────────────────────────────────────
def build_asset_registry(client_data: dict) -> str:
    raw_path   = build_asset_registry_docx(client_data)
    final_path = embed_fonts(raw_path)
    print("✅ สร้างบัญชีทรัพย์สินเสร็จ")
    return final_path
