"""
cover_builder.py (v3)
โครงสร้างตาม spec คุณพยัต:
  🔴 principle (hardcode)
  🟢 client_situation (AI gen — บางประเด็น)
  🔴 from_experience (hardcode)
  🔴 "ผมแนะนำให้:"
  🟢 bullets (AI gen — บางประเด็น)
  🔴 post_bullets / goal (hardcode)
"""
import os, json, tempfile
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Cordia New"
NAVY_RGB = RGBColor(0x1A, 0x3A, 0x5C)
DARK_RGB = RGBColor(0x22, 0x22, 0x22)
GRAY_RGB = RGBColor(0x55, 0x55, 0x55)
LIGHT_BLUE_HEX = "EBF2FA"
NAVY_HEX       = "1A3A5C"
YELLOW_HEX     = "FFFBEA"

MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-5-20250929")

# ─────────────────────────────────────────────────────────────────────
# FIXED CONTENT
# ─────────────────────────────────────────────────────────────────────

INTRO = (
    "การทำพินัยกรรมนั้น ถ้ามองแค่วิธีการจัดทำนั้นไม่ยากเลย "
    "เพียงแค่เขียนข้อความบอกว่าอยากให้ทรัพย์สินไปอยู่ที่ไหน "
    "ลงวันที่แล้วเซ็นชื่อต่อหน้าพยาน 2 คน เท่านี้ก็ถูกต้องตามกฎหมายทุกอย่าง "
    "ทายาทสามารถนำไปให้ศาลออกคำสั่งต่างๆ ได้แล้ว\n\n"
    "แต่สำหรับพ่อแม่ที่มีลูกเล็ก เรื่องนี้ต้องคิดให้ไกลกว่านั้น "
    "เพราะสิ่งที่คุณต้องการจริงๆ ไม่ใช่แค่เอกสารที่ถูกกฎหมาย "
    "แต่มันคือแผน แผนที่ทำให้คุณมั่นใจว่าถ้าวันหนึ่งเกิดเหตุไม่คาดฝัน "
    "คู่ชีวิตจะไม่ลำบาก ลูกๆจะยังมีคนดูแล ครอบครัวจะมีเงินพอใช้และจะได้รับสิ่งที่คุณตั้งใจให้จริงๆ\n\n"
    "เอกสารชุดนี้จึงครอบคลุม 12 ประเด็นที่พ่อแม่ลูกเล็กทุกคู่ควรพิจารณา ดังต่อไปนี้"
)

NEXT_STEPS = (
    "จากทั้ง 12 ประเด็นที่ผ่านมา "
    "จะเห็นได้ว่าการวางแผนไม่ได้จบแค่การมีพินัยกรรม\n\n"
    "แต่เป็นการออกแบบ \"คู่มือสำหรับเหตุการณ์ฉุกเฉิน\" "
    "เพื่อให้ครอบครัวสามารถใช้ชีวิตต่อได้จริง แม้ในวันที่เราไม่อยู่แล้ว\n\n"
    "ขั้นตอนถัดไปที่ผมแนะนำคือ:\n\n"
    "- สานต่อและเติมข้อมูลจริงของคุณลงในเอกสารต่างๆ ที่แนบมาพร้อมกัน ให้ครบถ้วนสมบูรณ์\n"
    "- พูดคุยกับบุคคลสำคัญ (คู่สมรส ผู้จัดการมรดก ผู้ปกครอง และผู้ดูแลเงิน) "
    "เพื่อให้ทุกคนเข้าใจบทบาทของตนเองอย่างชัดเจน\n"
    "- จัดทำพินัยกรรมให้ถูกต้องตามกฎหมาย นัดพยาน ลงนาม "
    "และจัดเก็บเอกสารให้เป็นระบบ พร้อมแจ้งผู้ที่เกี่ยวข้องทราบ\n"
    "- พิจารณาวางแผนประกันภัย เพื่อให้ความคุ้มครองสอดคล้องกับแผนที่วางไว้\n\n"
    "หากคุณต้องการให้ผมช่วยดูแลในส่วนของการวางแผนประกันภัย "
    "ผมยินดีช่วยออกแบบแผนความคุ้มครอง และแนะนำแนวทางที่เหมาะสมให้กับครอบครัวของคุณ\n\n"
    "ผมได้รับใบอนุญาต AFPT (Associate Financial Planner Thailand) "
    "และสามารถให้คำแนะนำด้านการวางแผนประกันภัยได้อย่างถูกต้องตามหลักวิชาชีพ "
    "โดยไม่ได้เป็นตัวแทนประกันภัย จึงสามารถให้ความเห็นได้อย่างเป็นกลาง\n\n"
    "ทั้งหมดนี้ เพื่อให้แผนฉบับนี้ ไม่ใช่แค่ \"เอกสารที่ดี\" "
    "แต่เป็น \"คู่มือของครอบครัวคุณที่ใช้ได้จริง\""
)

# ── ISSUE_FIXED ────────────────────────────────────────────────────────
# structure ต่อ issue:
#   title            : str
#   principle        : str   (RED ก่อน client_situation)
#   from_experience  : str|None  (RED ระหว่าง client_situation กับ bullets)
#   post_bullets     : str|None  (RED หลัง bullets เช่น bridge)
#   goal             : str|None  (RED goal)
#   has_situation    : bool  (🟢 client_situation)
#   has_ai_bullets   : bool  (🟢 bullets)
#   fixed_bullets    : list|None (🔴 fixed bullets สำหรับ 07, 11, 12)
#   fixed_situation  : str|None  (🔴 fixed situation สำหรับ 12)

ISSUE_FIXED = {
    1: {
        "title": "เงินพร้อมใช้ทันทีหลังจากไป",
        "principle": (
            "ทันทีที่เราจากไป โลกไม่ได้หยุดรอ ค่าใช้จ่ายจะเกิดขึ้นภายในไม่กี่ชั่วโมง "
            "ทั้งค่ารักษาพยาบาลครั้งสุดท้าย ค่าจัดงานศพ ค่าธรรมเนียมแจ้งตาย "
            "ไปจนถึงค่าดำเนินการทางกฎหมาย สิ่งเหล่านี้ต้องใช้ \"เงินสดทันที\"\n\n"
            "หากไม่ได้เตรียมไว้ คนที่เรารักต้องควักเงินของตัวเองออกไปก่อน "
            "ในช่วงเวลาที่เจ็บปวดที่สุดในชีวิต"
        ),
        "from_experience": (
            "จากประสบการณ์ของผม "
            "สิ่งสำคัญไม่ใช่แค่ \"มีเงินพอ\" แต่ต้อง \"หยิบใช้ได้ทันที\""
        ),
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้ครอบครัวสามารถ \"ใช้เงินได้ทันที\" "
            "โดยไม่ต้องติดขั้นตอน หรือรอการอนุมัติใดๆ"
        ),
        "has_situation": True, "has_ai_bullets": True,
        "fixed_bullets": None, "fixed_situation": None,
    },
    2: {
        "title": "เงินรองรับช่วงปรับตัว",
        "principle": (
            "การจากไปของคนในครอบครัว ไม่ได้กระทบแค่ความรู้สึก "
            "แต่กระทบ \"วิถีชีวิตทั้งหมด\" ในทันที "
            "คนที่อยู่ต่อ อาจต้องหยุดงานชั่วคราวเพื่อตั้งหลัก "
            "ดูแลลูก จัดการเอกสาร และปรับตัวกับสถานการณ์ใหม่ "
            "ขณะที่ค่าใช้จ่ายยังคงดำเนินต่อเหมือนเดิม\n\n"
            "ช่วง 6\u201312 เดือนแรก คือช่วงที่หนักที่สุด "
            "และจำเป็นต้องมี \"เงินก้อน\" เพื่อให้ครอบครัวสามารถตั้งหลักได้"
        ),
        "from_experience": (
            "จากประสบการณ์ของผม "
            "ครอบครัวที่ผ่านช่วงนี้ได้ดี คือครอบครัวที่ \"มีเวลา\" ให้ตัวเองได้ตั้งหลัก "
            "โดยไม่ต้องรีบตัดสินใจเรื่องสำคัญเพราะแรงกดดันทางการเงิน"
        ),
        "post_bullets": (
            "ในส่วนของรายได้ที่หายไปในระยะยาว "
            "จะเป็นอีกประเด็นสำคัญที่ต้องวางแผนเพิ่มเติม ซึ่งผมจะอธิบายในข้อถัดไป"
        ),
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้ครอบครัวมีเวลาในการปรับตัว "
            "โดยไม่ต้องตัดสินใจเรื่องสำคัญภายใต้ความเร่งรีบหรือความกดดันทางการเงิน"
        ),
        "has_situation": True, "has_ai_bullets": True,
        "fixed_bullets": None, "fixed_situation": None,
    },
    3: {
        "title": "รายได้ที่หายไปตลอดชีวิต",
        "principle": (
            "เมื่อเราจากไป สิ่งที่หายไปไม่ใช่แค่ตัวเรา "
            "แต่ \"รายได้\" ที่เราสามารถสร้างได้ในอนาคตก็หายไปด้วย"
        ),
        "from_experience": (
            "ประกันชีวิตจึงมีบทบาทสำคัญ "
            "ไม่ใช่ในฐานะเงินก้อน แต่เป็น \"เครื่องมือทดแทนรายได้\" "
            "เพื่อให้ครอบครัวยังสามารถใช้ชีวิตต่อได้ "
            "แม้ไม่มีผู้หารายได้หลักอยู่แล้ว"
        ),
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้ครอบครัวยังมี \"รายได้ต่อเนื่อง\" "
            "โดยไม่ต้องเปลี่ยนวิถีชีวิตทันทีหลังเกิดเหตุ"
        ),
        "has_situation": True, "has_ai_bullets": True,
        "fixed_bullets": None, "fixed_situation": None,
    },
    4: {
        "title": "ความต่อเนื่องของกิจการ",
        "principle": (
            "สำหรับครอบครัวที่มีธุรกิจของตัวเอง "
            "หากเจ้าของกิจการเกิดเหตุไม่คาดฝัน\n\n"
            "สิ่งที่มักเกิดขึ้นทันที ไม่ใช่แค่รายได้หายไป "
            "แต่คือ \"การตัดสินใจทั้งหมดหยุดลง\""
        ),
        "from_experience": (
            "จากประสบการณ์ของผม "
            "ความต่อเนื่องของกิจการ ขึ้นอยู่กับการที่ \"มีคนตัดสินใจแทนได้ทันที\""
        ),
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้กิจการสามารถ \"เดินต่อได้อย่างราบรื่น\" "
            "แม้ไม่มีเจ้าของอยู่ในช่วงเวลาหนึ่ง"
        ),
        "has_situation": True, "has_ai_bullets": True,
        "fixed_bullets": None, "fixed_situation": None,
    },
    5: {
        "title": "การพิการหรือทำงานไม่ได้",
        "principle": (
            "สิ่งที่น่ากลัวเหนือความตาย คือ "
            "การที่ยังมีชีวิตอยู่ แต่พิการหรือป่วยหนักจนไม่สามารถทำงานได้\n\n"
            "เพราะนอกจากจะไม่สามารถหารายได้ได้แล้ว "
            "ยังอาจกลายเป็นภาระให้กับครอบครัวในระยะยาว"
        ),
        "from_experience": (
            "จากประสบการณ์ของผม "
            "แม้เหตุการณ์ลักษณะนี้จะมีโอกาสเกิดขึ้นไม่บ่อย "
            "แต่หากเกิดขึ้นแล้ว จะเป็นสถานการณ์ที่กระทบครอบครัวมากที่สุด"
        ),
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้ครอบครัวยังมี \"เงินใช้ต่อเนื่อง\" "
            "และลดความเสี่ยงที่จะต้องกลายเป็นภาระของครอบครัวในระยะยาว"
        ),
        "has_situation": True, "has_ai_bullets": True,
        "fixed_bullets": None, "fixed_situation": None,
    },
    6: {
        "title": "ทรัพย์สินจะถึงมือลูกจริงไหม",
        "principle": (
            "หากทำพินัยกรรมโดยยกทรัพย์สินทั้งหมดให้คู่สมรสโดยไม่มีเงื่อนไข "
            "โดยหวังว่าเขาจะจัดการและส่งต่อให้ลูกอย่างยุติธรรม\n\n"
            "ในความเป็นจริง หากวันหนึ่งเกิดการเปลี่ยนแปลงในชีวิต "
            "เช่น การแต่งงานใหม่ ทรัพย์สินที่เราสร้างมาทั้งชีวิต "
            "อาจไม่ได้ถูกส่งต่อไปยังลูกตามที่ตั้งใจไว้ทั้งหมด\n\n"
            "เหตุการณ์ลักษณะนี้มีให้เห็นอยู่บ่อยครั้ง "
            "เพียงเพราะขาดการวางแผนล่วงหน้า\n\n"
            "ประเด็นนี้ไม่ใช่เรื่องของความไว้ใจ "
            "แต่คือการออกแบบให้ทรัพย์สิน \"ไปถึงปลายทางจริง\""
        ),
        "from_experience": None,  # ไม่มี — principle จบด้วย "ผมแนะนำให้:" โดยตรง
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้ทรัพย์สินสามารถ \"ไปถึงคนที่ตั้งใจจริง\" "
            "โดยไม่เปลี่ยนไปตามสถานการณ์ในอนาคต"
        ),
        "has_situation": True,
        "has_ai_bullets": True,
        "fixed_bullets": None, "fixed_situation": None,
    },
    7: {
        "title": "แผนสำรองหากคู่สมรสจากไปด้วย",
        "principle": (
            "พินัยกรรมส่วนใหญ่ มักเขียนขึ้นบนสมมติฐานว่า \"คู่สมรสยังมีชีวิตอยู่\" "
            "แต่ในความเป็นจริง อาจเกิดกรณีที่คู่สมรสจากไปก่อน "
            "หรือทั้งสองคนจากไปพร้อมกันได้\n\n"
            "หากไม่มีการวางแผนสำรองไว้ล่วงหน้า "
            "ทรัพย์มรดกอาจถูกแบ่งไปตามกฎหมายมรดก "
            "ซึ่งอาจไม่สอดคล้องกับความตั้งใจของครอบครัว"
        ),
        "from_experience": (
            "สถานการณ์นี้สามารถแก้ไขได้ด้วยการกำหนด \"แผนสำรอง\" ในพินัยกรรม"
        ),
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้ทรัพย์มรดกไม่ถูกแบ่งไปตามกฎหมายมรดก "
            "และยังคงส่งต่อไปยังลูกได้ตามที่ตั้งใจ "
            "แม้ในสถานการณ์ที่ไม่คาดคิดที่สุด"
        ),
        "has_situation": True,
        "has_ai_bullets": False,  # bullets เป็น RED fixed
        "fixed_bullets": [
            "- ระบุในพินัยกรรมว่า ในกรณีที่คู่สมรสถึงแก่ความตายก่อน "
            "หรือถึงแก่ความตายพร้อมกับข้าพเจ้า "
            "ให้ทรัพย์มรดกทั้งหมดตกเป็นของบุตรทั้งหมดของข้าพเจ้า"
        ],
        "fixed_situation": None,
    },
    8: {
        "title": "ใครดูแลลูกหากพ่อแม่จากไปพร้อมกัน",
        "principle": (
            "อุบัติเหตุที่พรากพ่อแม่ทั้งสองคนในคราวเดียว "
            "เป็นเหตุการณ์ที่ไม่มีใครอยากนึกถึง\n\n"
            "แต่หากเกิดขึ้นจริงและไม่ได้วางแผนไว้ล่วงหน้า "
            "ศาลจะเป็นผู้พิจารณาว่าเด็กควรอยู่กับใคร\n\n"
            "ในช่วงระหว่างการดำเนินการ อาจเกิดความไม่แน่นอนภายในครอบครัวได้ "
            "ทั้งในกรณีที่มีการโต้แย้งเรื่องการดูแลเด็กเมื่อมีประเด็นทรัพย์มรดก "
            "หรือในทางกลับกัน ไม่มีผู้ที่เตรียมตัวหรือพร้อมรับผิดชอบอย่างชัดเจน"
        ),
        "from_experience": (
            "จากประสบการณ์ของผม "
            "เรื่องที่กระทบเด็กมากที่สุด ไม่ใช่แค่ \"จะไปอยู่กับใคร\" "
            "แต่คือ \"ความชัดเจนและความพร้อมของคนที่ดูแล\""
        ),
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้ลูกมี \"คนดูแลที่ชัดเจน\" "
            "และลดความไม่แน่นอนในช่วงเวลาที่สำคัญที่สุดของชีวิต"
        ),
        "has_situation": True, "has_ai_bullets": True,
        "fixed_bullets": None, "fixed_situation": None,
    },
    9: {
        "title": "คนดูแลลูก กับ คนดูแลเงินของลูก",
        "principle": (
            "หากผู้ปกครองและผู้ดูแลเงินของลูกเป็นคนเดียวกัน "
            "จะไม่มีกลไกในการตรวจสอบว่าเงินถูกใช้เพื่อประโยชน์ของลูกอย่างเหมาะสมหรือไม่\n\n"
            "การแยกสองบทบาทนี้ออกจากกัน ไม่ได้หมายถึงความไม่ไว้ใจ "
            "แต่เป็นการออกแบบระบบให้เกิด \"การตรวจสอบและถ่วงดุล\" ระหว่างกัน"
        ),
        "from_experience": (
            "จากประสบการณ์ของผม "
            "การแยกบทบาทระหว่าง \"คนดูแลชีวิต\" และ \"คนดูแลเงิน\" "
            "เป็นสิ่งสำคัญในการปกป้องผลประโยชน์ของลูกในระยะยาว"
        ),
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้ทรัพย์สินของลูกได้รับการดูแลอย่างเหมาะสม "
            "และมีระบบตรวจสอบที่ชัดเจนในระยะยาว"
        ),
        "has_situation": True, "has_ai_bullets": True,
        "fixed_bullets": None, "fixed_situation": None,
    },
    10: {
        "title": "คู่มือฉุกเฉินสำหรับคนที่รับช่วงต่อ",
        "principle": (
            "เมื่อเกิดเหตุไม่คาดฝันขึ้น ทุกอย่างมักเกิดขึ้นพร้อมกันในช่วงเวลาสั้นๆ\n\n"
            "คนที่ต้องรับหน้าที่จัดการเรื่องทั้งหมด "
            "จำเป็นต้องรู้ว่าเอกสารอยู่ที่ไหน "
            "ต้องเริ่มจากอะไร และต้องดำเนินการเรื่องใดก่อนหลัง\n\n"
            "ทั้งในด้านการจัดการเอกสาร การติดต่อหน่วยงานต่างๆ "
            "และการใช้สิทธิ์ที่เกี่ยวข้องภายในระยะเวลาที่กำหนด\n\n"
            "หากไม่มีการเตรียม \"คู่มือ\" ไว้ล่วงหน้า "
            "ความล่าช้าอาจทำให้ครอบครัวเสียสิทธิ์ "
            "หรือได้รับเงินช้ากว่าที่ควร"
        ),
        "from_experience": (
            "จากประสบการณ์ของผม "
            "แผนที่ดี ไม่ใช่แค่ \"คิดไว้\" แต่ต้อง \"ทำให้คนอื่นทำต่อได้\""
        ),
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ คือทำให้คนที่รับช่วงต่อสามารถ \"ลงมือทำได้ทันที\" "
            "โดยไม่ต้องเสียเวลาเริ่มต้นใหม่ในช่วงเวลาที่สำคัญที่สุด"
        ),
        "has_situation": True,
        "has_ai_bullets": False,  # bullets เป็น RED fixed
        "fixed_bullets": [
            "- จัดทำคู่มือขั้นตอนหลังเกิดเหตุให้เป็นส่วนหนึ่งของเอกสารชุดนี้ "
            "(ผมจะจัดทำโครงร่างไว้ให้ และแนะนำให้ปรับรายละเอียดเพิ่มเติมให้สมบูรณ์ตามข้อมูลจริง)",
            "- จัดระบบเอกสารให้ชัดเจน และให้ผู้จัดการมรดกเข้าถึงได้ทันทีเมื่อจำเป็น",
            "- ทบทวนและอัปเดตข้อมูลอย่างน้อยปีละครั้ง หรือเมื่อมีการเปลี่ยนแปลงสำคัญ",
        ],
        "fixed_situation": None,
    },
    11: {
        "title": "ทบทวนแผนทุกปี",
        "principle": (
            "ชีวิตเปลี่ยนตลอดเวลา "
            "อาจมีลูกเพิ่มขึ้น รายได้เพิ่มหรือลดลง "
            "ผู้ปกครองที่เลือกไว้อาจเสียชีวิตก่อน "
            "หรือสถานการณ์อื่นๆ ที่เปลี่ยนแปลงไปจากที่เราคาดไม่ถึง\n\n"
            "แผนที่ดี คือแผนที่เราปรับมันให้ทันต่อสถานการณ์ปัจจุบันอยู่ตลอด"
        ),
        "from_experience": (
            "จากประสบการณ์ของผม "
            "พินัยกรรมที่ไม่เคยทบทวน "
            "มักเป็นพินัยกรรมที่ \"สร้างปัญหาแทนที่แก้ไขปัญหา\" ในวันที่สำคัญที่สุด"
        ),
        "post_bullets": None,
        "goal": (
            "เป้าหมายของส่วนนี้ "
            "คือให้แผนครอบครัวของคุณ \"เดินตามทันชีวิตจริง\" "
            "ไม่ใช่เอกสารที่หมดอายุก่อนเวลา"
        ),
        "has_situation": True, "has_ai_bullets": True,
        "fixed_bullets": None, "fixed_situation": None,
    },
    12: {
        "title": "จดหมายถึงครอบครัว",
        "principle": (
            "ไม่มีเอกสารกฎหมายฉบับใด สามารถทดแทน \"จดหมายจากคนที่เรารัก\" ได้\n\n"
            "จดหมายนี้ไม่ใช่เรื่องของกฎหมาย "
            "แต่เป็นพื้นที่สำหรับถ่ายทอดสิ่งที่ไม่สามารถใส่ไว้ในพินัยกรรมได้ "
            "ไม่ว่าจะเป็นความรัก เหตุผลเบื้องหลังการตัดสินใจ "
            "ค่านิยมที่อยากส่งต่อ หรือสิ่งที่อยากฝากไว้ให้ลูกใช้เป็นแนวทางในการใช้ชีวิต"
        ),
        "from_experience": (
            "จากประสบการณ์ของผม "
            "จดหมายฉบับนี้ มักเป็นสิ่งที่ครอบครัวจดจำและมีคุณค่ามากที่สุดในระยะยาว"
        ),
        "post_bullets": None,
        "goal": None,  # goal อยู่ใน fixed_situation block แล้ว
        "has_situation": True,   # AI gen ใช้ชื่อคู่สมรสและจำนวนลูกจริง
        "has_ai_bullets": False,
        "fixed_bullets": [
            "- พิจารณาเขียนจดหมายถึงคู่สมรสและลูกแต่ละคน แยกเป็นรายบุคคล",
            "- ถ่ายทอดความตั้งใจ เหตุผลของการวางแผนทรัพย์สิน และสิ่งที่อยากฝากไว้",
            "- จัดเก็บจดหมายไว้ร่วมกับเอกสารพินัยกรรม "
            "และมอบหมายให้ผู้จัดการมรดกเป็นผู้ส่งมอบในช่วงเวลาที่เหมาะสม",
        ],
        "fixed_situation": None,
    },
}


# ─────────────────────────────────────────────────────────────────────
# Claude: generate client_situation + bullets
# ─────────────────────────────────────────────────────────────────────

def _call_issues(client, client_data: dict, issue_nums: list, include_summary: bool) -> dict:
    """Call Claude สำหรับ issue_nums ที่ระบุ"""

    INSTRUCTIONS = {
        1: """   ประเด็น 1: situation 2 ย่อหน้า — (1) เงินสดมีเท่าไหร่ เพียงพอกับ funeral_budget+40,000 บาทไหม (2) สวัสดิการมีอะไรบ้าง แต่ใช้เวลานาน นับเป็นเงินฉุกเฉินไม่ได้
   ประเด็น 1: bullets 3 ข้อ — (1) บัญชีชื่อเดียวแนะนำแจ้ง ATM แยกบัญชีฉุกเฉิน (2) สวัสดิการใช้เวลา ไม่ควรนับระยะสั้น (3) ระบุในพินัยกรรมให้ผู้จัดการมรดกเบิกเงินได้ทันที""",
        2: """   ประเด็น 2: situation 1 ย่อหน้า — หนี้เท่าไหร่ รายได้ที่หายไป ค่าใช้จ่ายปรับตัว 12 เดือน = 70% ของรายได้ ระบุตัวเลขจริง
   ประเด็น 2: bullets 3 ข้อ — (1) คำนวณ debt + income×70%×12 บอกว่ายังขาดเท่าไหร่ (2) ประกันชีวิตระบุผู้รับเป็นคู่สมรสโดยตรง (3) รวมกับข้อ 3 ได้""",
        3: """   ประเด็น 3: situation 2 ย่อหน้า — (1) อายุ รายได้ HLV=PV(4%,60-อายุ,รายได้×12) ครอบครัวพึ่ง 100% (2) สรุปต้องมีเครื่องมือทดแทนรายได้
   ประเด็น 3: bullets 3 ข้อ — (1) คำนวณทุน = debt+70%×12+HLV-ทรัพย์สินสภาพคล่อง-ประกันที่มี แนะนำขั้นต่ำ 50% (2) ผู้รับเป็นคู่สมรสโดยตรง (3) Term ความยาวอิงลูกคนเล็ก+20 ปี""",
        4: """   ประเด็น 4: ถ้าไม่มีกิจการ → situation "ไม่มีกิจการ ไม่เกี่ยวข้อง" bullets ["- ประเด็นนี้อยู่ในเกณฑ์ดีแล้ว ไม่มีสิ่งที่ต้องดำเนินการเพิ่มเติม"]
   ถ้ามีกิจการ → situation ระบุประเภท รายได้ พนักงาน แผนรองรับ / bullets 3 ข้อ: (1) หนังสือมอบอำนาจให้คู่สมรสหรือผู้จัดการ (2) เงินสำรองกิจการ 6 เดือน ระบุตัวเลข (3) Keyman Insurance""",
        5: """   ประเด็น 5: situation — มีประกันโรคร้าย/ทุพพลภาพอะไรบ้าง ผลกระทบ 2 กรณี (โรคร้าย + ทุพพลภาพ)
   ประเด็น 5: bullets 3 ข้อ — (1) ทุพพลภาพวงเงินเท่าทุนข้อ 3 ระบุตัวเลข (2) โรคร้าย income×12 ระบุตัวเลข (3) แนบท้ายประกันหลัก""",
        6: """   ประเด็น 6: situation — กังวลเรื่องนี้ไหม ทรัพย์สินหลักมีอะไร ชื่อโฉนดใคร ถ้ายกทุกอย่างให้คู่สมรสโดยไม่มีเงื่อนไขจะเกิดอะไร
   ประเด็น 6: bullets ตามทรัพย์สินจริง — บ้าน→ยกลูกแต่คู่สมรสอยู่อาศัย ประกัน→แบ่งผู้รับ หุ้น/ลงทุน→ระบุในพินัยกรรม""",
        7: """   ประเด็น 7: situation — แผนแบ่งทรัพย์เป็นอย่างไร ถ้าไม่มีแผนสำรองและตายพร้อมกัน กฎหมายไทยจะแบ่งอย่างไร (รวมพ่อแม่ของลูกค้า) / bullets ส่งกลับ []""",
        8: """   ประเด็น 8: situation — ตั้งใจให้ใครดูแลลูก ได้คุยแล้วหรือยัง มีสำรองไหม
   ประเด็น 8: bullets 3 ข้อ — (1) คุยผู้ปกครองหลักก่อนลงนาม (2) หาสำรอง คุยให้เข้าใจหน้าที่ (3) ระบุในพินัยกรรมแยกจากผู้ดูแลเงิน""",
        9: """   ประเด็น 9: situation — มีแนวคิดแยกผู้ปกครองกับผู้ดูแลเงินหรือไม่ ให้ใคร คุยกันแล้วหรือยัง มีสำรองไหม
   ประเด็น 9: bullets 3 ข้อ — (1) แยกหน้าที่ชัดเจนในพินัยกรรม (2) เตรียมสำรองอย่างน้อย 1 คน (3) พูดคุยให้เข้าใจว่ามีหน้าที่ตามกฎหมายเพื่อประโยชน์เด็กเท่านั้น""",
        10: """   ประเด็น 10: situation — เอกสารเก็บที่ไหน ใครรู้บ้าง มีคู่มือหลังเกิดเหตุไหม ถ้าเกิดเหตุวันนี้จะเริ่มทำอะไรก่อน / bullets ส่งกลับ []""",
        11: """   ประเด็น 11: situation — อายุลูกค้า คู่สมรส ลูกแต่ละคน แผนนี้เป็นแผนแรก ส่วนที่ทำได้ทันที vs ส่วนที่ต้องติดตามต่อ (ประกัน)
   ประเด็น 11: bullets 3 ข้อ — (1) จัดทำตามลำดับ แม้ไม่ครบก็ดีกว่าไม่มี ตั้งเตือน calendar ปีละครั้ง (2) ติดตามประกันชีวิต โรคร้าย ทุพพลภาพ (3) ทุกครั้งที่ชีวิตเปลี่ยนให้ปรับแผน""",
        12: """   ประเด็น 12: situation — ขึ้นต้น "ปัจจุบัน คุณ[ชื่อ]ยังไม่ได้เขียนจดหมายถึง[ชื่อคู่สมรส]และลูกๆ..." ระบุชื่อคู่สมรสและลูกทุกคน / bullets ส่งกลับ []""",
    }

    instructions = "\n\n".join(INSTRUCTIONS[n] for n in issue_nums if n in INSTRUCTIONS)
    nums_str = ", ".join(str(n) for n in issue_nums)

    summary_block = ""
    if include_summary:
        summary_block = """
นอกจากนี้ให้ gen "action_summary" โดยดูจากสถานะ 12 ประเด็น จัดกลุ่ม 3 ระดับ:
กลุ่ม 1 "ทำได้ทันที" — ไม่ต้องใช้เงิน ลงมือได้เลย
กลุ่ม 2 "ต้องดำเนินการ" — ใช้เวลาแต่ไม่ซับซ้อน
กลุ่ม 3 "วางแผนเพิ่มเติม" — ต้องใช้เงินหรือปรึกษาผู้เชี่ยวชาญ

output format:
{{"issues": [...], "action_summary": {{"ทำได้ทันที": ["- ..."], "ต้องดำเนินการ": ["- ..."], "วางแผนเพิ่มเติม": ["- ..."]}}}}"""
    else:
        summary_block = 'output format:\n{{"issues": [...]}}'

    prompt = f"""คุณคือผู้ช่วยของคุณพยัต นักวางแผนการเงินและกฎหมายในประเทศไทย

ข้อมูลลูกค้า:
{json.dumps(client_data, ensure_ascii=False, indent=2)}

งาน: gen สถานการณ์ลูกค้า + คำแนะนำ สำหรับประเด็น {nums_str} เท่านั้น

instruction ต่อประเด็น:
{instructions}

กฎ:
- ขึ้นต้น situation ด้วย "ปัจจุบัน คุณ[ชื่อ]..."
- ใส่ "คุณ" นำหน้าชื่อเสมอ
- ภาษาไทยล้วน ยกเว้น Keyman Insurance, AFPT, Term
- โทนมืออาชีพแต่อบอุ่น
- ตัวเลขต้องสอดคล้องกัน

ข้อบังคับ: ตอบด้วย JSON เท่านั้น ห้ามมีข้อความอื่นใดก่อนหรือหลัง JSON

{summary_block}

ตัวอย่าง format issues:
[{{"ลำดับ": {issue_nums[0]}, "สถานการณ์ลูกค้า": "ปัจจุบัน คุณ...", "คำแนะนำ": ["- ...", "- ..."]}}]"""

    resp = client.messages.create(
        model=MODEL, max_tokens=8000,
        messages=[{"role": "user", "content": prompt}]
    )
    text = resp.content[0].text.replace("```json","").replace("```","").strip()
    print(f"DEBUG batch {nums_str}: {text[:100]}...")

    decoder = json.JSONDecoder()
    # หา [ สำหรับ issues array หรือ { สำหรับ object
    for start_char in ['{', '[']:
        idx = text.find(start_char)
        if idx != -1:
            try:
                obj, _ = decoder.raw_decode(text, idx)
                if isinstance(obj, list):
                    return {"issues": obj, "action_summary": {}}
                if isinstance(obj, dict):
                    return obj
            except json.JSONDecodeError:
                pass
    print(f"⚠️ batch {nums_str}: ไม่พบ JSON ที่ valid")
    return {"issues": [], "action_summary": {}}


def generate_issue_content(client_data: dict, issues: list = None) -> dict:
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    print("กำลัง generate ประเด็น 1-6...")
    batch1 = _call_issues(client, client_data, list(range(1, 7)), include_summary=False)

    print("กำลัง generate ประเด็น 7-12 + action_summary...")
    batch2 = _call_issues(client, client_data, list(range(7, 13)), include_summary=True)

    all_issues = batch1.get("issues", []) + batch2.get("issues", [])
    action_summary = batch2.get("action_summary", {})

    print(f"✅ generate เสร็จ: {len(all_issues)} ประเด็น")

    result = {item["ลำดับ"]: item for item in all_issues if isinstance(item, dict) and "ลำดับ" in item}
    result["__action_summary__"] = action_summary
    return result

# ─────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────

def _set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color); shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def _r(para, text, bold=False, sz=15, color=None, italic=False):
    run = para.add_run(text)
    run.font.name = FONT; run.font.size = Pt(sz)
    run.bold = bold; run.italic = italic
    run.font.color.rgb = color or DARK_RGB
    return run

def _p(doc, text, sz=15, bold=False, color=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=8, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(sz * 1.5)
    p.alignment = align
    _r(p, text, bold=bold, sz=sz, color=color, italic=italic)
    return p

def _multi_para(doc, text, sz=15, color=None, sb_first=6, sa=4):
    """แยก \n\n เป็นหลาย paragraph"""
    parts = [t.strip() for t in text.split("\n\n") if t.strip()]
    for i, part in enumerate(parts):
        _p(doc, part, sz=sz, color=color, sb=sb_first if i == 0 else 2, sa=sa)


# ─────────────────────────────────────────────────────────────────────
# Build docx
# ─────────────────────────────────────────────────────────────────────

def build_cover(client_data: dict, issues: list, generated: dict, folder_name: str) -> str:
    nickname = client_data.get("ชื่อเล่น", folder_name.split("_")[0])
    date_str = folder_name.split("_")[1] if "_" in folder_name else ""

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.79)
        sec.bottom_margin = Inches(0.79)
        sec.left_margin   = Inches(0.79)
        sec.right_margin  = Inches(0.79)

    # ── Header ────────────────────────────────────────────────────────
    htbl = doc.add_table(rows=1, cols=1); htbl.style = "Table Grid"
    hcell = htbl.cell(0, 0); _set_bg(hcell, NAVY_HEX)
    p1 = hcell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(8); p1.paragraph_format.space_after = Pt(4)
    _r(p1, "แผนคุ้มครองครอบครัว", bold=True, sz=28, color=RGBColor(0xFF, 0xFF, 0xFF))
    p2 = hcell.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    _r(p2, f"คุณ{nickname}", sz=18, color=RGBColor(0xBB, 0xDD, 0xFF))
    p3 = hcell.add_paragraph(); p3.paragraph_format.space_after = Pt(8)
    _r(p3, f"จัดทำโดย คุณพยัต  |  {date_str[:2]}/{date_str[2:4]}/{date_str[4:]}",
       sz=13, color=RGBColor(0x88, 0xAA, 0xCC))
    doc.add_paragraph()

    # ── Intro ─────────────────────────────────────────────────────────
    _multi_para(doc, INTRO, sz=15, color=DARK_RGB, sb_first=0, sa=10)
    doc.add_paragraph()

    # ── 12 Issues ─────────────────────────────────────────────────────
    for num in range(1, 13):
        f   = ISSUE_FIXED[num]
        gen = generated.get(num, {})
        ai_situation = gen.get("สถานการณ์ลูกค้า", "")
        ai_bullets   = gen.get("คำแนะนำ", [])

        # Issue header
        tbl = doc.add_table(rows=1, cols=2); tbl.style = "Table Grid"
        tbl.columns[0].width = Inches(0.5); tbl.columns[1].width = Inches(6.5)
        nc = tbl.cell(0, 0); _set_bg(nc, NAVY_HEX)
        np = nc.paragraphs[0]; np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _r(np, f"{num:02d}", bold=True, sz=16, color=RGBColor(0xFF, 0xFF, 0xFF))
        tc = tbl.cell(0, 1); _set_bg(tc, LIGHT_BLUE_HEX)
        tp = tc.paragraphs[0]
        _r(tp, f["title"], bold=True, sz=16, color=NAVY_RGB)

        # 🔴 Principle
        _multi_para(doc, f["principle"], sz=15, color=DARK_RGB, sb_first=6, sa=4)

        # 🟢 Client situation (AI gen หรือ fixed_situation)
        if f["has_situation"] and ai_situation:
            _p(doc, ai_situation, sz=15, color=DARK_RGB, sb=4, sa=4)
        elif f.get("fixed_situation"):
            sit = f["fixed_situation"].format(n=nickname)
            _p(doc, sit, sz=15, color=DARK_RGB, sb=4, sa=4)

        # 🔴 From experience + "ผมแนะนำให้:"
        if f.get("from_experience"):
            _p(doc, f["from_experience"], sz=15, color=DARK_RGB, sb=4, sa=4)

        # "ผมแนะนำให้:" label (สำหรับประเด็นที่มี bullets)
        has_any_bullets = f["has_ai_bullets"] or f.get("fixed_bullets")
        if has_any_bullets:
            _p(doc, "ผมแนะนำให้:", sz=15, bold=False, color=DARK_RGB, sb=2, sa=2)

        # 🟢 AI bullets
        if f["has_ai_bullets"] and ai_bullets:
            for b in ai_bullets:
                _p(doc, b, sz=15, color=DARK_RGB, sb=0, sa=3)

        # 🔴 Fixed bullets
        if f.get("fixed_bullets"):
            for b in f["fixed_bullets"]:
                _p(doc, b, sz=15, color=DARK_RGB, sb=0, sa=3)

        # 🔴 Post bullets (bridge)
        if f.get("post_bullets"):
            _p(doc, f["post_bullets"], sz=14, color=GRAY_RGB, sb=4, sa=4, italic=False)

        # 🔴 Goal
        if f.get("goal"):
            _p(doc, f["goal"], sz=15, color=NAVY_RGB, sb=6, sa=14)


    # ── สรุปสิ่งที่ต้องทำ ─────────────────────────────────────
    action_summary = generated.get("__action_summary__", {})
    if action_summary:
        doc.add_paragraph()
        atbl = doc.add_table(rows=1, cols=1); atbl.style = "Table Grid"
        acell = atbl.cell(0, 0); _set_bg(acell, NAVY_HEX)
        ap = acell.paragraphs[0]
        ap.paragraph_format.space_before = Pt(6); ap.paragraph_format.space_after = Pt(4)
        _r(ap, "สรุปสิ่งที่ต้องทำ", bold=True, sz=18, color=RGBColor(0xFF,0xFF,0xFF))

        GROUP_COLORS = {
            "ทำได้ทันที": RGBColor(0x1A, 0x7A, 0x4A),
            "ต้องดำเนินการ": RGBColor(0x1A, 0x3A, 0x5C),
            "วางแผนเพิ่มเติม": RGBColor(0xAA, 0x00, 0x00),
        }
        GROUP_BG = {
            "ทำได้ทันที": "E8F5EE",
            "ต้องดำเนินการ": "EBF2FA",
            "วางแผนเพิ่มเติม": "FFF0F0",
        }
        for group_name in ["ทำได้ทันที", "ต้องดำเนินการ", "วางแผนเพิ่มเติม"]:
            items = action_summary.get(group_name, [])
            if not items: continue
            # header
            htbl2 = doc.add_table(rows=1, cols=1); htbl2.style = "Table Grid"
            hcell2 = htbl2.cell(0, 0); _set_bg(hcell2, GROUP_BG[group_name])
            hp = hcell2.paragraphs[0]
            hp.paragraph_format.space_before = Pt(6); hp.paragraph_format.space_after = Pt(2)
            _r(hp, group_name, bold=True, sz=14, color=GROUP_COLORS[group_name])
            for item in items:
                _p(doc, item, sz=14, color=DARK_RGB, sb=0, sa=3)

    # ── ขั้นตอนถัดไป ──────────────────────────────────────────────────
    doc.add_paragraph()
    ntbl = doc.add_table(rows=1, cols=1); ntbl.style = "Table Grid"
    ncell = ntbl.cell(0, 0); _set_bg(ncell, LIGHT_BLUE_HEX)
    np2 = ncell.paragraphs[0]
    np2.paragraph_format.space_before = Pt(6); np2.paragraph_format.space_after = Pt(4)
    _r(np2, "ขั้นตอนถัดไป", bold=True, sz=18, color=NAVY_RGB)
    _multi_para(doc, NEXT_STEPS, sz=15, color=DARK_RGB, sb_first=8, sa=4)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    print("✅ Built cover page")
    return tmp.name
