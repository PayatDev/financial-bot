"""
emergency_guide_builder.py — doc 5
คู่มือ 90 วันแรก: fixed structure + fill ชื่อจริงจาก JSON
"""
import os, re, tempfile, zipfile, shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT     = "Cordia New"
BLACK    = RGBColor(0x00, 0x00, 0x00)
GRAY     = RGBColor(0x88, 0x88, 0x88)
RED      = RGBColor(0xAA, 0x00, 0x00)
NAVY     = RGBColor(0x1A, 0x3A, 0x5C)
AMBER    = RGBColor(0x7A, 0x58, 0x00)
AMBER_BG = "FFFBEA"
NAVY_BG  = "EBF2FA"
LS       = 1.3
FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")

_NA = {'ยังไม่มี', '-', 'ยังไม่ได้กำหนด', 'ไม่ได้ระบุ', 'ไม่มี', ''}


# ── data helpers ──────────────────────────────────────────────────────
def _get(d, *keys):
    for k in keys:
        v = str(d.get(k, '') or '').strip()
        if v and v not in _NA:
            return v
    return ''

def _fill(d, placeholder, *keys):
    v = _get(d, *keys)
    return v if v else f"[{placeholder}]"

def _clean(s):
    return re.sub(r'\s*\(.*?\)', '', s or '').strip()


# ── text helpers ──────────────────────────────────────────────────────
def _run(para, text, bold=False, sz=13, color=None, underline=False):
    r = para.add_run(text)
    r.font.name = FONT; r.font.size = Pt(sz)
    r.bold = bold; r.font.color.rgb = color or BLACK
    if underline: r.underline = True
    return r

def _add_runs(para, text, sz=13):
    for part in re.split(r'(\[.*?\])', text):
        if not part: continue
        if part.startswith('['):
            _run(para, part, bold=True, sz=sz, color=RED, underline=True)
        else:
            _run(para, part, sz=sz)

def _p(doc, text, sz=13, bold=False, color=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    para.alignment = align
    if bold or color:
        _run(para, text, bold=bold, sz=sz, color=color)
    else:
        _add_runs(para, text, sz=sz)

def _section(doc, title, sz=15, sb=12, sa=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    _run(para, title, bold=True, sz=sz, color=NAVY)

def _subsection(doc, title, sz=13, sb=8, sa=2):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    _run(para, title, bold=True, sz=sz)

def _note_box(doc, text, sz=11):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), AMBER_BG); shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
    tcBdr = OxmlElement('w:tcBdr')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:color'),'D4A000'); tcBdr.append(el)
    tcPr.append(tcBdr)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after  = Pt(4)
    cp.paragraph_format.line_spacing = Pt(sz * LS)
    _run(cp, text, sz=sz, color=AMBER)


# ── table helpers ─────────────────────────────────────────────────────
def _set_col_width(table, col_idx, width_pct, total_cm=15.5):
    """set column width in cm"""
    from docx.shared import Cm
    for row in table.rows:
        row.cells[col_idx].width = Cm(total_cm * width_pct)

def _tbl_header(tbl, headers, sz=11):
    row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
        r = cell.paragraphs[0].add_run(h)
        r.font.name = FONT; r.font.size = Pt(sz)
        r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        # fill navy
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1E3A5F"); shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

def _tbl_row(tbl, row_idx, cells, sz=11, bg=None):
    row = tbl.rows[row_idx]
    for i, txt in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = Pt(sz * LS)
        # ใช้ _add_runs เพื่อ render placeholder สีแดง
        for part in re.split(r'(\[.*?\])', txt):
            if not part: continue
            r = p.add_run(part)
            r.font.name = FONT; r.font.size = Pt(sz)
            if part.startswith('['):
                r.bold = True; r.font.color.rgb = RED; r.underline = True
            else:
                r.font.color.rgb = BLACK
        if bg:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), bg); shd.set(qn("w:val"), "clear")
            tcPr.append(shd)

def _timeline_table(doc, header_text, rows_data, col_widths=(0.35, 0.18, 0.25, 0.22)):
    """
    rows_data: list of (งาน, ผู้รับผิดชอบ, ติดต่อ, เอกสาร)
    """
    n_rows = len(rows_data) + 1  # +1 header
    tbl = doc.add_table(rows=n_rows, cols=4)
    tbl.style = "Table Grid"
    _tbl_header(tbl, ["งาน", "ผู้รับผิดชอบ", "ติดต่อ/หน่วยงาน", "เอกสารที่ต้องเตรียม"])
    for i, row_data in enumerate(rows_data):
        bg = "F5F7FA" if i % 2 == 0 else "FFFFFF"
        _tbl_row(tbl, i+1, row_data, bg=bg)

def _access_table(doc, headers, rows_data, col_widths=None):
    """ตารางสำหรับหมวดข้อมูลการเข้าถึง"""
    n_rows = len(rows_data) + 1
    tbl = doc.add_table(rows=n_rows, cols=len(headers))
    tbl.style = "Table Grid"
    _tbl_header(tbl, headers)
    for i, row_data in enumerate(rows_data):
        bg = "F5F7FA" if i % 2 == 0 else "FFFFFF"
        _tbl_row(tbl, i+1, row_data, bg=bg)


# ── build docx ────────────────────────────────────────────────────────
def build_guide_docx(client_data: dict) -> str:
    d  = client_data
    n  = _get(d, "ชื่อเล่น", "nickname") or "ผู้ทำเอกสาร"
    sp = _clean(_get(d, "คู่สมรส", "spouse_nickname")) or "คู่สมรส"
    occ = _get(d, "อาชีพ", "occupation")
    crypto = _get(d, "คริปโต", "assets_crypto_wallet")

    executor   = _fill(d, "ชื่อผู้จัดการมรดก",   "ผู้จัดการมรดก", "estate_executor")
    urgent_mgr = _fill(d, "ชื่อผู้จัดการฉุกเฉิน", "ผู้จัดการฉุกเฉิน", "urgent_manager")
    doc_loc    = _fill(d, "ที่เก็บเอกสาร",        "ที่อยู่เอกสาร", "documents_location")
    funeral    = _get(d, "ความปรารถนางานศพ", "funeral_wishes") or "ตามที่ระบุไว้ในพินัยกรรม"
    insurance  = _get(d, "ประกันชีวิต", "insurance_life")
    welfare    = _get(d, "สวัสดิการ", "welfare")
    debt       = _get(d, "หนี้สิน", "debt")
    property_  = _get(d, "อสังหาริมทรัพย์", "assets_property")
    executor_name = _fill(d, "ชื่อ-นามสกุลผู้จัดการมรดก",
                          "ชื่อ-นามสกุลจริงผู้จัดการมรดก", "fullname_executor")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.787)
        section.bottom_margin = Inches(0.787)
        section.left_margin   = Inches(1.181)
        section.right_margin  = Inches(0.787)

    # ── title ─────────────────────────────────────────────────────────
    _p(doc, f"คู่มือ 90 วันแรกของครอบครัวคุณ{n}",
       sz=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    _p(doc, "เอกสารนี้เป็นร่าง — กรอกข้อมูลจริงและเก็บในที่ปลอดภัยก่อนใช้",
       sz=11, color=RED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)

    # ── ข้อมูลด่วน ────────────────────────────────────────────────────
    _note_box(doc,
        f"⚠️  ติดต่อก่อนเลย:\n"
        f"ผู้จัดการมรดก: {executor_name}  โทร [เบอร์ผู้จัดการมรดก]\n"
        f"ผู้จัดการฉุกเฉิน: {urgent_mgr}  โทร [เบอร์ผู้จัดการฉุกเฉิน]\n"
        f"เอกสารทั้งหมดอยู่ที่: {doc_loc}\n"
        f"ทนาย: คุณพยัต จิรสุวรรณพงศ์  payat.jira@gmail.com")

    # ── ช่วงที่ 1 — 7 วันแรก ─────────────────────────────────────────
    _section(doc, "ช่วงที่ ๑ — 7 วันแรก")

    rows_7 = [
        ("แจ้งตายที่สำนักงานเขต/อำเภอ ภายใน 24 ชั่วโมง",
         executor_name,
         "สำนักงานเขต/อำเภอใกล้บ้าน",
         "ใบรับรองแพทย์ + บัตรประชาชนผู้ตาย"),
        (f"จัดงานศพ — {funeral}",
         executor_name,
         "[วัดหรือสถานที่จัดงาน]",
         "ใบมรณบัตร"),
        ("แจ้งธนาคารทุกแห่งที่มีบัญชี ขออายัดหรือดูแลบัญชีชั่วคราว",
         executor_name,
         "ธนาคารทุกแห่งที่มีบัญชี",
         "ใบมรณบัตร + บัตรประชาชนผู้จัดการมรดก"),
        (f"แจ้งที่ทำงาน/ต้นสังกัด{(' — ' + occ) if occ else ''}",
         executor_name,
         "[ชื่อหน่วยงาน] ฝ่าย HR",
         "ใบมรณบัตร"),
        (f"เปิดพินัยกรรมและแจ้งผู้จัดการมรดก — เอกสารอยู่ที่ {doc_loc}",
         executor_name,
         "คุณพยัต จิรสุวรรณพงศ์",
         "พินัยกรรม"),
    ]
    _timeline_table(doc, "7 วันแรก", rows_7)

    # ── ช่วงที่ 2 — 30 วัน ───────────────────────────────────────────
    _section(doc, "ช่วงที่ ๒ — ภายใน 30 วัน")

    rows_30 = []
    # ประกันชีวิต
    ins_detail = insurance if insurance else "[ชื่อบริษัทประกัน]"
    rows_30.append((
        "แจ้งบริษัทประกันชีวิตทุกกรมธรรม์ — ยื่นเรียกร้องค่าสินไหม",
        executor_name,
        ins_detail,
        "ใบมรณบัตร + กรมธรรม์ + บัตรประชาชนผู้รับประโยชน์"
    ))

    # สวัสดิการ
    has_welfare = welfare and welfare not in _NA
    if has_welfare:
        rows_30.append((
            f"ยื่นขอรับสิทธิสวัสดิการ — {welfare}",
            executor_name,
            "หน่วยงานที่เกี่ยวข้อง",
            "ใบมรณบัตร + บัตรสมาชิก/บัตรประชาชน"
        ))
    else:
        rows_30.append((
            "ยื่นขอรับสิทธิสวัสดิการ (กบข./บำเหน็จ/ประกันสังคม ถ้ามี)",
            executor_name,
            "กบข. 02-636-1000 | สปส. 1506",
            "ใบมรณบัตร + บัตรสมาชิก"
        ))

    rows_30.append((
        "รวบรวมเอกสารยื่นขอตั้งผู้จัดการมรดกต่อศาล",
        executor_name,
        "ศาลแพ่ง 02-512-8345 | คุณพยัต จิรสุวรรณพงศ์",
        "พินัยกรรม + ทะเบียนบ้าน + บัตรประชาชนผู้จัดการ"
    ))
    _timeline_table(doc, "30 วัน", rows_30)

    # ── ช่วงที่ 3 — 90 วัน ───────────────────────────────────────────
    _section(doc, "ช่วงที่ ๓ — ภายใน 90 วัน")

    rows_90 = [
        ("ดำเนินการตามพินัยกรรม — แบ่งทรัพย์สินตามที่ระบุไว้",
         executor_name,
         "คุณพยัต จิรสุวรรณพงศ์",
         "พินัยกรรม + คำสั่งศาลแต่งตั้งผู้จัดการมรดก"),
    ]

    if property_ and property_ not in _NA:
        rows_90.append((
            f"โอนกรรมสิทธิ์อสังหาริมทรัพย์ — {property_}",
            executor_name,
            "กรมที่ดิน 02-141-5555",
            "โฉนด + คำสั่งศาล + บัตรประชาชนทุกฝ่าย"
        ))
    else:
        rows_90.append((
            "โอนกรรมสิทธิ์อสังหาริมทรัพย์ (ถ้ามี)",
            executor_name,
            "กรมที่ดิน 02-141-5555",
            "โฉนด + คำสั่งศาล + บัตรประชาชนทุกฝ่าย"
        ))

    if debt and debt not in _NA:
        rows_90.append((
            f"จัดการหนี้สิน — {debt}",
            executor_name,
            "เจ้าหนี้แต่ละราย",
            "สัญญากู้ + ใบมรณบัตร + คำสั่งศาล"
        ))
    else:
        rows_90.append((
            "จัดการหนี้สินที่เหลือ (ถ้ามี)",
            executor_name,
            "เจ้าหนี้แต่ละราย",
            "สัญญากู้ + ใบมรณบัตร"
        ))

    rows_90.append((
        "อัปเดตผู้รับประโยชน์ในเอกสารทุกฉบับ — ประกัน กองทุน บัญชีธนาคาร",
        f"[คู่สมรส/ลูก]",
        "บริษัทประกัน / ธนาคาร / กองทุน",
        "คำสั่งศาล + กรมธรรม์/บัญชีเดิม"
    ))
    _timeline_table(doc, "90 วัน", rows_90)

    # ── หมวดข้อมูลการเข้าถึง (กรอกมือ) ─────────────────────────────
    _section(doc, "หมวดพิเศษ — ข้อมูลการเข้าถึงบัญชีสำคัญ", sb=16)
    _note_box(doc,
        "⚠️  กรอกด้วยลายมือบนกระดาษ ใส่ซองปิดผนึก เก็บในเซฟ\n"
        "ห้ามพิมพ์ ถ่ายรูป หรือส่งทางออนไลน์เด็ดขาด")

    _subsection(doc, "อีเมลหลัก", sb=8)
    _access_table(doc,
        ["รายการ", "ข้อมูล"],
        [
            ["Email address", "________________________"],
            ["Password", "________________________"],
            ["วิธี 2FA (เช่น SMS / Authenticator app)", "________________________"],
            ["Recovery email / phone", "________________________"],
        ]
    )

    _subsection(doc, "Platform / Social / บัญชีสำคัญ", sb=10)
    _access_table(doc,
        ["Platform", "Username / Email", "Password", "วิธีเข้าถึง 2FA"],
        [
            ["________________________", "________________________", "________________________", "________________________"],
            ["________________________", "________________________", "________________________", "________________________"],
            ["________________________", "________________________", "________________________", "________________________"],
        ]
    )

    _subsection(doc, "App ธนาคาร / Mobile Banking", sb=10)
    _access_table(doc,
        ["ธนาคาร", "เลขบัญชี", "PIN app", "วิธี reset / เบอร์ธนาคาร"],
        [
            ["________________________", "________________________", "________________________", "________________________"],
            ["________________________", "________________________", "________________________", "________________________"],
            ["________________________", "________________________", "________________________", "________________________"],
        ]
    )

    _subsection(doc, f"Crypto Wallet{' — มี' if crypto else ' — (กรอกถ้ามี)'}", sb=10)
    _access_table(doc,
        ["รายการ", "ข้อมูล"],
        [
            ["Wallet / Exchange", "________________________"],
            ["Seed phrase เก็บที่", "________________________"],
            ["วิธีเข้าถึง / PIN", "________________________"],
            ["⚠️ ห้ามเปิดเผย Seed phrase ทางออนไลน์", "ปรึกษา claude.ai ก่อนดำเนินการใดๆ"],
        ]
    )

    _subsection(doc, "รหัสอุปกรณ์", sb=10)
    _access_table(doc,
        ["อุปกรณ์", "PIN / Password", "ที่เก็บ Backup"],
        [
            ["โทรศัพท์มือถือ", "________________________", "________________________"],
            ["คอมพิวเตอร์", "________________________", "________________________"],
            ["Password Manager (เช่น Bitwarden)", "Master password: ________________________", "________________________"],
        ]
    )

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name); tmp.close()
    return tmp.name


# ── embed fonts ───────────────────────────────────────────────────────
def embed_fonts(docx_path: str) -> str:
    out = docx_path.replace(".docx", "_emb.docx")
    shutil.copy(docx_path, out)
    font_map = {
        "THSarabunNew":            os.path.join(FONT_DIR, "THSarabunNew.ttf"),
        "THSarabunNew-Bold":       os.path.join(FONT_DIR, "THSarabunNew Bold.ttf"),
        "THSarabunNew-Italic":     os.path.join(FONT_DIR, "THSarabunNew Italic.ttf"),
        "THSarabunNew-BoldItalic": os.path.join(FONT_DIR, "THSarabunNew BoldItalic.ttf"),
    }
    with zipfile.ZipFile(out, "a") as z:
        for name, path in font_map.items():
            arcname = f"word/fonts/{name}.ttf"
            try: z.getinfo(arcname)
            except KeyError:
                if os.path.exists(path): z.write(path, arcname)
    with zipfile.ZipFile(out, "r") as z:
        names = z.namelist()
        ft = z.read("word/fontTable.xml").decode()
        ct = z.read("[Content_Types].xml").decode()
        fr = (z.read("word/_rels/fontTable.xml.rels").decode()
              if "word/_rels/fontTable.xml.rels" in names else
              '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
              '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>')
        all_files = {n: z.read(n) for n in names}
    old = '<w:font w:name="TH Sarabun New"/>'
    new = ('<w:font w:name="TH Sarabun New">'
           '<w:embedRegular r:id="rIdF1"/><w:embedBold r:id="rIdF2"/>'
           '<w:embedItalic r:id="rIdF3"/><w:embedBoldItalic r:id="rIdF4"/></w:font>')
    ft = ft.replace(old, new) if old in ft else ft.replace("</w:fonts>", new + "</w:fonts>")
    rels = ''.join([
        '<Relationship Id="rIdF1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew.ttf"/>',
        '<Relationship Id="rIdF2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Bold.ttf"/>',
        '<Relationship Id="rIdF3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Italic.ttf"/>',
        '<Relationship Id="rIdF4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-BoldItalic.ttf"/>',
    ])
    fr = fr.replace("</Relationships>", rels + "</Relationships>")
    if "ttf" not in ct:
        ct = ct.replace("</Types>", '<Default Extension="ttf" ContentType="application/x-font-ttf"/></Types>')
    all_files["word/fontTable.xml"]            = ft.encode()
    all_files["word/_rels/fontTable.xml.rels"] = fr.encode()
    all_files["[Content_Types].xml"]            = ct.encode()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in all_files.items(): z.writestr(name, data)
    os.unlink(docx_path)
    return out


# ── entry point ───────────────────────────────────────────────────────
def build_emergency_guide(client_data: dict) -> str:
    raw_path   = build_guide_docx(client_data)
    final_path = embed_fonts(raw_path)
    print("✅ สร้างคู่มือ 90 วันเสร็จ")
    return final_path
