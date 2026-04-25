"""
poa_builder.py
สร้างหนังสือมอบอำนาจ Word 2 หน้า
- หน้า 1: วิธีใช้ (fixed)
- หน้า 2: หนังสือมอบอำนาจ (fill จาก client_data)
"""
import os
import re
import tempfile
import zipfile
import shutil

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT     = "TH Sarabun New"
BLACK    = RGBColor(0x00, 0x00, 0x00)
GRAY     = RGBColor(0x88, 0x88, 0x88)
RED      = RGBColor(0xAA, 0x00, 0x00)
LS       = 1.3
FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")


# ── Text helpers ──────────────────────────────────────────────────────
def _run(para, text, bold=False, sz=13, color=None, underline=False):
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(sz)
    run.bold = bold
    run.font.color.rgb = color or BLACK
    if underline:
        run.underline = True
    return run

def _add_runs(para, text, sz=13):
    """[placeholder] → สีแดง underline"""
    for part in re.split(r'(\[.*?\])', text):
        if not part:
            continue
        if part.startswith('['):
            _run(para, part, bold=True, sz=sz, color=RED, underline=True)
        else:
            _run(para, part, sz=sz)

def _p(doc, text, sz=13, bold=False, color=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    para.alignment = align
    if bold or color:
        _run(para, text, bold=bold, sz=sz, color=color)
    else:
        _add_runs(para, text, sz=sz)

def _head(doc, text, sz=14, sb=8, sa=2):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    _run(para, text, bold=True, sz=sz)

def _bullet(doc, text, sz=13):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    para.paragraph_format.left_indent  = Pt(12)
    _run(para, "- ", sz=sz)
    _add_runs(para, text, sz=sz)

def _clause(doc, n, text, sz=13):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.line_spacing = Pt(sz * LS)
    _run(para, f"{n}.  ", bold=True, sz=sz)
    _add_runs(para, text, sz=sz)


# ── Signature table ───────────────────────────────────────────────────
def _no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:color'), 'auto')
        tcBdr.append(el)
    tcPr.append(tcBdr)

def _no_table_border(tbl):
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "auto")
        tblBdr.append(el)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBdr)

def _no_table_border_OLD(tbl):
    """ลบ border ระดับ table ทั้งหมด"""
    tblPr = tbl._tbl.get_or_add_tblPr()
    tblBdr = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:color'), 'auto')
        tblBdr.append(el)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBdr)

def _sig_table(doc, pairs, sz=12, sb=8):
    """
    pairs = [(role, name_ph), ...]  1 หรือ 2 คน
    กึ่งกลาง, มีที่ว่างให้เซ็น, ขอบใส, เส้นสีเทาอ่อน
    """
    n_cols   = len(pairs)
    TOTAL    = 9000   # DXA
    GAP      = 800
    col_w    = (TOTAL - GAP * (n_cols - 1)) // n_cols
    col_widths = []
    for i in range(n_cols):
        col_widths.append(col_w)
        if i < n_cols - 1:
            col_widths.append(GAP)

    n_tc = len(col_widths)
    tbl = doc.add_table(rows=3, cols=n_tc)
    tbl.style = "Table Grid"
    _no_table_border(tbl)

    for i, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[i].width = w * 635

    for row in tbl.rows:
        for cell in row.cells:
            _no_border(cell)

    tbl.rows[0].height = Pt(30)
    tbl.rows[1].height = Pt(2)
    tbl.rows[2].height = Pt(sz * LS + 2)

    content_cols = [i * 2 for i in range(n_cols)]

    # row 0 — ที่ว่างเซ็น + role label ขวาล่าง
    for idx, (role, _) in enumerate(pairs):
        col = content_cols[idx]
        cell = tbl.rows[0].cells[col]
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cp.paragraph_format.space_before = Pt(sb if idx == 0 else 0)
        cp.paragraph_format.space_after  = Pt(0)
        cp.paragraph_format.line_spacing = Pt(sz * LS)
        _run(cp, role, sz=sz, color=GRAY)

    # row 1 — เส้นใต้ลายเซ็น (top border)
    for idx in range(n_cols):
        col = content_cols[idx]
        cell = tbl.rows[1].cells[col]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBdr = OxmlElement('w:tcBdr')
        for side in ['top', 'left', 'bottom', 'right']:
            el = OxmlElement(f'w:{side}')
            if side == 'top':
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '6')
                el.set(qn('w:color'), 'AAAAAA')
            else:
                el.set(qn('w:val'), 'none')
                el.set(qn('w:sz'), '0')
                el.set(qn('w:color'), 'auto')
            tcBdr.append(el)
        old = tcPr.find(qn('w:tcBdr'))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(tcBdr)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # row 2 — ชื่อ กึ่งกลาง
    for idx, (_, name_ph) in enumerate(pairs):
        col = content_cols[idx]
        cell = tbl.rows[2].cells[col]
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(6)
        cp.paragraph_format.line_spacing = Pt(sz * LS)
        _run(cp, "(", sz=sz - 1, color=GRAY)
        _add_runs(cp, name_ph, sz=sz - 1)
        _run(cp, ")", sz=sz - 1, color=GRAY)


# ── Build docx ────────────────────────────────────────────────────────
def build_poa_docx(client_data: dict) -> str:
    d = client_data
    n  = d.get("ชื่อเล่น", "ผู้ทำพินัยกรรม")
    sp = d.get("คู่สมรส",  "คู่สมรส")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.787)   # 2 cm
        section.bottom_margin = Inches(0.787)
        section.left_margin   = Inches(1.181)   # 3 cm
        section.right_margin  = Inches(0.787)

    # ══════════════════════════════════════════════════
    # PAGE 1 — วิธีใช้ (fixed)
    # ══════════════════════════════════════════════════
    _p(doc, "วิธีใช้หนังสือมอบอำนาจฉบับนี้",
       sz=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)

    _head(doc, "ใช้เมื่อไหร่", sz=14, sb=4, sa=2)
    _bullet(doc, "ผู้มอบอำนาจอยู่ในโรงพยาบาลหรือไม่สามารถเดินทางได้")
    _bullet(doc, "ผู้มอบอำนาจต้องการให้คนอื่นทำธุรกรรมแทนโดยไม่ต้องมาเอง")

    _head(doc, "หน่วยงานที่ใช้ได้", sz=14, sb=8, sa=2)
    _bullet(doc, "ธนาคารทุกแห่ง — ฝาก ถอน โอน ปิดบัญชี และจัดการบัญชีหลักทรัพย์")
    _bullet(doc, "สำนักงานประกันสังคม — เรียกร้องสิทธิ์และรับเงินแทน")
    _bullet(doc, "บริษัทประกันภัย — เรียกร้องค่าสินไหมและจัดการกรมธรรม์")
    _bullet(doc, "กรมสรรพากร — ยื่นแบบและรับเงินคืนภาษีแทน")
    _bullet(doc, "หน่วยงานราชการทั่วไป — ติดต่อและยื่นเอกสารแทน")

    _head(doc, "ใช้ไม่ได้เมื่อไหร่", sz=14, sb=8, sa=2)
    _bullet(doc, "หลังผู้มอบอำนาจเสียชีวิต — หนังสือฉบับนี้สิ้นผลทันที ให้ใช้พินัยกรรมแทน")
    _bullet(doc, "กรมที่ดิน — ต้องใช้แบบฟอร์ม ท.ด.21 ของกรมที่ดินโดยเฉพาะ")
    _bullet(doc, "ธนาคารบางแห่งอาจมีแบบฟอร์มของตัวเอง อาจต้องกรอกเพิ่มเติม")

    _head(doc, "วิธีใช้งาน", sz=14, sb=8, sa=2)
    _bullet(doc, "ใช้ต้นฉบับเท่านั้น ห้ามใช้สำเนา")
    _bullet(doc, "พกบัตรประชาชนของผู้รับมอบอำนาจไปด้วยทุกครั้ง")
    _bullet(doc, "พกสำเนาบัตรประชาชนของผู้มอบอำนาจไปด้วย")

    _head(doc, "จำนวนชุดที่ควรเตรียม", sz=14, sb=8, sa=2)
    _p(doc,
       "ควรเตรียมอย่างน้อย 4 ชุด ได้แก่ ธนาคาร บริษัทหลักทรัพย์ บริษัทประกัน "
       "และสำรองเก็บที่บ้าน 1 ชุด โดยต้องเซ็นชื่อต้นฉบับทุกฉบับพร้อมกัน"
       "ต่อหน้าพยานในคราวเดียว",
       sa=0)

    # ══════════════════════════════════════════════════
    # PAGE 2 — หนังสือมอบอำนาจ
    # ══════════════════════════════════════════════════
    doc.add_page_break()

    _p(doc, "หนังสือมอบอำนาจ",
       sz=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    _p(doc, "ฉบับร่าง — โปรดตรวจสอบและลบบรรทัดนี้ก่อนพิมพ์จริง",
       sz=11, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, sa=8)

    _p(doc, "ทำที่ [ที่อยู่จัดทำเอกสาร]", sa=2)
    _p(doc, "วันที่ [วันที่ทำหนังสือ] เดือน [เดือน] พ.ศ. [ปี]", sa=6)

    _head(doc, "ผู้มอบอำนาจ", sb=4, sa=2)
    _p(doc,
       f"ข้าพเจ้า [ชื่อ-นามสกุลคุณ{n}] อายุ [อายุคุณ{n}] ปี สัญชาติไทย "
       f"ถือบัตรประจำตัวประชาชนเลขที่ [เลขบัตรประชาชนคุณ{n}] "
       f"อยู่บ้านเลขที่ [ที่อยู่ปัจจุบันคุณ{n}] "
       "ซึ่งต่อไปในหนังสือนี้จะเรียกว่า ผู้มอบอำนาจ")

    _head(doc, "ผู้รับมอบอำนาจ", sb=4, sa=2)
    _p(doc,
       f"[ชื่อ-นามสกุลคุณ{sp}] อายุ [อายุคุณ{sp}] ปี สัญชาติไทย "
       f"ถือบัตรประจำตัวประชาชนเลขที่ [เลขบัตรประชาชนคุณ{sp}] "
       f"อยู่บ้านเลขที่ [ที่อยู่ปัจจุบันคุณ{sp}] "
       "ซึ่งต่อไปในหนังสือนี้จะเรียกว่า ผู้รับมอบอำนาจ")

    _head(doc, "ขอบเขตของการมอบอำนาจ", sb=4, sa=2)
    _p(doc,
       "โดยหนังสือฉบับนี้ ผู้มอบอำนาจขอมอบอำนาจให้ผู้รับมอบอำนาจ"
       "มีอำนาจกระทำการแทนในกิจการดังต่อไปนี้", sa=2)

    scope_items = [
        "ดำเนินการเกี่ยวกับบัญชีธนาคารทุกบัญชีที่มีในชื่อผู้มอบอำนาจ "
        "รวมถึงการฝาก ถอน โอน ปิดบัญชี และขอเปิดบัญชีใหม่",
        "ดำเนินการเกี่ยวกับหุ้นและหลักทรัพย์ทุกประเภท "
        "รวมถึงการซื้อ ขาย โอน และปิดบัญชีหลักทรัพย์",
        "ดำเนินการเกี่ยวกับกองทุนสำรองเลี้ยงชีพของผู้มอบอำนาจ "
        "รวมถึงการรับผลประโยชน์ที่เกี่ยวข้อง",
        "ดำเนินการเกี่ยวกับสัญญาประกันภัยทุกฉบับ "
        "รวมถึงการเรียกร้องค่าสินไหม การต่ออายุ และการยกเลิกกรมธรรม์",
        "ดำเนินการติดต่อสำนักงานประกันสังคม "
        "รวมถึงการเรียกร้องสิทธิ์และรับเงินแทนผู้มอบอำนาจ",
        "ดำเนินการติดต่อกรมสรรพากร "
        "รวมถึงการยื่นแบบและรับเงินคืนภาษีแทนผู้มอบอำนาจ",
        "ลงนามในเอกสารทางกฎหมายและทางการเงินที่เกี่ยวข้องกับทรัพย์สินของผู้มอบอำนาจ",
        "ดำเนินการติดต่อกับหน่วยงานราชการและสถาบันการเงินในนามของผู้มอบอำนาจ",
    ]
    for num, text in zip(["๑","๒","๓","๔","๕","๖","๗","๘"], scope_items):
        _clause(doc, num, text)

    _head(doc, "เงื่อนไขการมีผล", sb=4, sa=2)
    _p(doc,
       "หนังสือฉบับนี้มีผลทันทีนับแต่วันลงนาม และยังคงมีผลแม้ผู้มอบอำนาจ"
       "จะไม่สามารถสื่อสารหรือตัดสินใจได้ตามปกติ "
       "หนังสือฉบับนี้สิ้นผลทันทีเมื่อผู้มอบอำนาจถึงแก่ความตาย "
       "และผู้มอบอำนาจสามารถเพิกถอนได้ทุกเมื่อโดยทำเป็นหนังสือแจ้ง "
       "การกระทำใดๆ ของผู้รับมอบอำนาจภายในขอบเขตข้างต้น "
       "ให้ถือว่าเป็นการกระทำของผู้มอบอำนาจทุกประการ", sa=6)

    _p(doc, "เพื่อเป็นหลักฐาน คู่สัญญาทั้งสองฝ่ายได้ลงลายมือชื่อไว้ต่อหน้าพยาน", sa=2)

    _sig_table(doc, [
        ("ผู้มอบอำนาจ",    f"ชื่อ-นามสกุลจริงคุณ{n}"),
        ("ผู้รับมอบอำนาจ", f"ชื่อ-นามสกุลจริงคุณ{sp}"),
    ], sz=12, sb=6)

    _p(doc, "พยานรับรองว่าคู่สัญญาทั้งสองฝ่ายได้ลงนามต่อหน้าพยานจริง",
       sb=8, sa=2)

    _sig_table(doc, [
        ("พยานที่ ๑", "ชื่อ-นามสกุลพยานที่ ๑"),
        ("พยานที่ ๒", "ชื่อ-นามสกุลพยานที่ ๒"),
    ], sz=12, sb=4)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


# ── Embed fonts ───────────────────────────────────────────────────────
def embed_fonts(docx_path: str) -> str:
    out = docx_path.replace(".docx", "_emb.docx")
    shutil.copy(docx_path, out)

    font_map = {
        "THSarabunNew":            os.path.join(FONT_DIR, "THSarabunNew.ttf"),
        "THSarabunNew-Bold":       os.path.join(FONT_DIR, "THSarabunNew Bold.ttf"),
        "THSarabunNew-Italic":     os.path.join(FONT_DIR, "THSarabunNew Italic.ttf"),
        "THSarabunNew-BoldItalic": os.path.join(FONT_DIR, "THSarabunNew BoldItalic.ttf"),
    }
    with zipfile.ZipFile(out, "a") as z:
        for name, path in font_map.items():
            arcname = f"word/fonts/{name}.ttf"
            try:
                z.getinfo(arcname)
            except KeyError:
                if os.path.exists(path):
                    z.write(path, arcname)

    with zipfile.ZipFile(out, "r") as z:
        names = z.namelist()
        ft = z.read("word/fontTable.xml").decode()
        ct = z.read("[Content_Types].xml").decode()
        fr = z.read("word/_rels/fontTable.xml.rels").decode() \
             if "word/_rels/fontTable.xml.rels" in names else \
             '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>' \
             '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'
        all_files = {n: z.read(n) for n in names}

    old = '<w:font w:name="TH Sarabun New"/>'
    new = (
        '<w:font w:name="TH Sarabun New">'
        '<w:embedRegular r:id="rIdF1"/>'
        '<w:embedBold r:id="rIdF2"/>'
        '<w:embedItalic r:id="rIdF3"/>'
        '<w:embedBoldItalic r:id="rIdF4"/>'
        "</w:font>"
    )
    ft = ft.replace(old, new) if old in ft else ft.replace("</w:fonts>", new + "</w:fonts>")

    rels = "".join([
        '<Relationship Id="rIdF1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew.ttf"/>',
        '<Relationship Id="rIdF2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Bold.ttf"/>',
        '<Relationship Id="rIdF3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-Italic.ttf"/>',
        '<Relationship Id="rIdF4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="fonts/THSarabunNew-BoldItalic.ttf"/>',
    ])
    fr = fr.replace("</Relationships>", rels + "</Relationships>")
    if "ttf" not in ct:
        ct = ct.replace("</Types>", '<Default Extension="ttf" ContentType="application/x-font-ttf"/></Types>')

    all_files["word/fontTable.xml"] = ft.encode()
    all_files["word/_rels/fontTable.xml.rels"] = fr.encode()
    all_files["[Content_Types].xml"] = ct.encode()

    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in all_files.items():
            z.writestr(name, data)

    os.unlink(docx_path)
    return out


# ── Entry point ───────────────────────────────────────────────────────
def build_poa(client_data: dict) -> str:
    """คืน path ไฟล์ .docx พร้อมส่งลูกค้า"""
    raw_path   = build_poa_docx(client_data)
    final_path = embed_fonts(raw_path)
    print("✅ สร้างหนังสือมอบอำนาจเสร็จ")
    return final_path
